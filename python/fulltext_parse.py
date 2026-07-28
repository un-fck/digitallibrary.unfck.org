#!/usr/bin/env python3
"""Semantic parser for UN resolution / decision / PRST full texts (Track A, v1).

Reads the low-interpretation extraction layer (digitallibrary.document_paragraphs_raw)
and emits ONE semantic JSON per document to
    <ARCHIVE_ROOT>/parsed_dev/<sanitized_symbol>.json

It classifies each raw paragraph into semantic elements (frontmatter, title,
heading, preambular/operative paragraph, footnote, vote record, signature,
table, divider) and enforces an accounting invariant: every raw position is
consumed by exactly one element's `positions[]` or appears in the top-level
`dropped[]` list with a reason.

Three format families are handled (see fulltext_census.py / the corpus census):
  - native `docx`  : rich styles, italic preambular lead runs (props.lead_italic_text),
                     literal operative numbers ("1.\\t..."), real kind='footnote' rows,
                     masthead often in table cells.
  - legacy `doc`   : usually styled like native; masthead frequently in table cells.
  - `wpd` (WP)     : ONE 'Normal' style, no italics, no numbering. Hierarchy is carried
                     by literal tabs / indent_firstline and text patterns; preambular
                     clauses are plain "Verb ... ," lines; footnotes are inline "N/ text"
                     after a "____" divider; sentences are sometimes hard-broken across
                     consecutive paragraph rows (we merge them).

This is deliberately standalone (mirrors the other python/ scripts): DATABASE_URL
from .env, short-lived psycopg (v3) connections. Targets are documents whose
document_files.status is 'extracted' or 'parsed' (so a re-parse still finds docs
already loaded to the semantic DB).

DB MODE (--to-db). The frozen semantic layer lands in two tables added by
migration 003:
  - digitallibrary.document_paragraphs — one row per parsed element (the JSON
    `elements[]`), document order. `id` is uuid5(NAMESPACE_URL,
    '<symbol_normalized>:<lang>:<position>') where position is the 0-based element
    index, computed by this loader.
  - digitallibrary.document_parses     — one row per (symbol,lang): parser_version,
    format, element_count, and the JSON root `dropped[]`/`issues[]` verbatim, so
    the accounting invariant stays queryable in SQL.
Loading is DELETE-then-INSERT per (symbol,lang) in BOTH tables (idempotent /
re-parsable), batched over short-lived connections of ~20 docs each, mirroring
fulltext_extract_raw.py's discipline. On success the document_files status is
advanced 'extracted' -> 'parsed'; on a hard parse/insert failure it is set to
'parse_failed' with the error recorded (never crashes the batch). An accounting
failure does NOT fail the load — the doc is still written and the failure is
recorded in document_parses.issues (and document_paragraphs stays queryable).
JSON output is written alongside the DB rows unless --db-only is given.

Usage:
    uv run python python/fulltext_parse.py                 # JSON only (all extracted/parsed docs)
    uv run python python/fulltext_parse.py --limit 5
    uv run python python/fulltext_parse.py --symbol A/RES/48/70
    uv run python python/fulltext_parse.py --to-db         # JSON + semantic DB
    uv run python python/fulltext_parse.py --db-only        # semantic DB only, no JSON
"""

from __future__ import annotations

import argparse
import json
import re
import uuid
from pathlib import Path

import fulltext_verbs as fv
from fulltext_common import ARCHIVE_ROOT, get_conn, sanitize_symbol, upsert_document_file

# sem-v2 adds the action-verb annotation pass (migration 004): a nested `action`
# object on each operative/preambular element, flattened to the action_*/assignee_*
# columns by the loader. Element construction / accounting are unchanged.
# sem-v3 (HEADING STRUCTURE): (1) STYLE-TRUSTED headings — a raw heading style
# (H1/H2/H23/H4../HCh/Heading1-9) becomes type='heading' with its style tier,
# with the TOC verifier's front-matter FALSE-POSITIVE exclusions ported; (2) a
# bare section-marker heading ('I') merges with the short title line after it
# (prefix='I.', text='General guidelines'); (3) whole-paragraph-bold structural
# labels ('Action 13.', 'Goal 1.', 'Objective 5:') + short bold title fragments
# become headings nested below their section; (4) leading markers emit as `prefix`;
# (5) annex/appendix delimiters carry subtype + label prefix + merged title, and
# every annexed element inherits its delimiter's annex_index. paragraph_type is
# untouched; text/positions/accounting are preserved (markers move to `prefix`,
# which the text-preservation gate counts).
# sem-v5 (MARKER SEQUENCE CONFIRMATION + SOURCE CONSERVATION): (1) a leading 'N.'
# is only read as an operative marker when it CONTINUES THE DOCUMENT'S OWN sequence
# (see detect_op_numbers) — the extraction layer merges page numbers and footnote
# references onto continuation fragments, which used to be stored as invented
# operative numbers (S/RES/661(1990) held a '19.' on a mid-sentence fragment of
# operative 3(c)); an unconfirmed number now stays inside the text and the element
# gets no prefix and no operative label. (2) the same confirmation applies to
# upper-case Roman parentheticals, which are otherwise the tail of a resolution
# citation ('resolution 900 (IX) of 11 December 1954'). (3) every parse is checked
# against its ARCHIVED SOURCE FILE (words in the source, or in the target document's
# span of a multi-document source page, vs words the parse holds) so a document that
# arrives truncated FAILS instead of being stored as 'parsed' — the defect that left
# 145 documents at ~21% of their source with no flag at all.
# sem-v4 (DOTTED ENUMERATORS + LONG SECTION HEADINGS): (1) a paragraph opening with a
# CONFIRMED dotted enumerator ('1.1 By 2030 ...', '8.10 ...', '1.a ...') — the SDG-
# target / forest-goal / programme-plan family — moves its marker to `prefix`, gets
# level=2, and has its text stripped; confirmation is sequence-based (a contiguous
# follows()-chain >=3 opening at a numeric '.1'), so assessment-scale rows ('0.01
# Albania ....'), statistics tables and lone in-sentence decimals are never promoted;
# (2) a bold/heading-styled lettered-or-roman section heading with a possibly-LONG
# descriptive title ('B. Advancing integration ...', 'III. Green economy ...') is now
# recognised even past the 80-char guard, with a verb guard so a bold 'I. Decides ...'
# operative is not swallowed. paragraph_type/positions/accounting are untouched.
PARSER_VERSION = "sem-v5"
OUT_DIR = ARCHIVE_ROOT / "parsed_dev"

# Opt-in source-defect rescue: when a resolution drops an operative NUMBER at
# source (e.g. A/HRC/RES/19/1 omits '5.'/'6.'; S/RES/1528(2004) omits '(e)'-'(h)'),
# an unlabeled clause that (a) sits inside a running operative sequence with a
# CONFIRMED numbering gap ahead and (b) reads like an operative (finite lead verb
# or a 'To <verb>' sub-item) is labeled operative with props.inferred_operative and
# NO invented prefix. Default ON; recorded per element so it is auditable/reversible.
RESCUE_INFERRED_OPERATIVE = True

# 'To <lowercase-verb>' opens an infinitive operative sub-item (…(a) To observe,
# (b) To liaise…); used by the rescue to recognise an unlabeled sibling sub-item.
INFINITIVE_SUBITEM_RE = re.compile(r"^To\s+[a-z]")

# ---------------------------------------------------------------------------
# Vocabularies & patterns
# ---------------------------------------------------------------------------

# Preambular lead verbs (first word, case-insensitive). Used for the WP fallback
# where there is no props.lead_italic_text to read the verb from. Deliberately
# broad; the state machine (opening-formula .. first-operative) is the primary
# signal, this only labels lead_verb and rescues stray clauses.
PREAMBULAR_FIRST_WORDS = {
    "recalling", "reaffirming", "noting", "recognizing", "recognising",
    "welcoming", "considering", "convinced", "concerned", "emphasizing",
    "emphasising", "expressing", "guided", "having", "bearing", "taking",
    "mindful", "alarmed", "acknowledging", "determined", "determining",
    "stressing", "underlining", "underscoring", "desiring", "desirous",
    "aware", "regretting", "deploring", "affirming", "observing", "believing",
    "conscious", "deeply", "gravely", "firmly", "fully", "further",
    "reiterating", "recalling", "seeking", "encouraged", "endorsing",
    "commending", "appreciating", "anxious", "cognizant", "cognisant",
    "remaining", "keeping", "realizing", "realising", "supporting",
    "welcoming", "invoking", "conscious", "noting", "having", "aware",
    "highlighting", "acknowledging", "reaffirming", "expressing", "guided",
    "concerned", "convinced", "confident", "hopeful", "resolved", "eager",
    "grateful", "pleased", "sharing", "aiming",
}

# Second words that extend a lead-verb phrase (e.g. "Recalling also",
# "Deeply concerned", "Taking note", "Bearing in mind", "Having considered").
PREAMBULAR_SECOND_WORDS = {
    "also", "further", "note", "in", "with", "that", "the", "concerned",
    "convinced", "aware", "mindful", "considered", "examined", "recalled",
    "regard", "account", "into", "of", "deeply", "again",
}

# Finite operative lead verbs (3rd-person present). Distinct from the participial
# preambular leads -- used to catch UNNUMBERED operative clauses in short
# resolutions ("Decides to hold ...") that carry no "1." prefix.
OPERATIVE_LEAD_VERBS = {
    "decides", "requests", "calls", "urges", "invites", "recommends",
    "reaffirms", "notes", "welcomes", "encourages", "endorses", "expresses",
    "takes", "considers", "demands", "stresses", "emphasizes", "emphasises",
    "reiterates", "approves", "adopts", "condemns", "deplores", "declares",
    "affirms", "acknowledges", "appeals", "authorizes", "authorises",
    "designates", "proclaims", "resolves", "supports", "commends", "confirms",
    "underlines", "underscores", "recognizes", "recognises", "determines",
    "establishes", "requires", "instructs", "directs",
    "implores", "pledges", "undertakes",
}

# Opening formula. Tolerates a leading quote (A/HRC/PRST statements quote the
# Council resolution verbatim: '"The Human Rights Council,') so those PRSTs are
# recognised as resolution-structured (preambular/operative labeling applies).
OPENING_RE = re.compile(
    r"^[\"“”‘’']?\s*The\s+(General Assembly|Security Council|Economic and Social Council|"
    r"Human Rights Council|Trusteeship Council)\s*,?\s*$"
)

# Title patterns.
TITLE_GA_NUM_RE = re.compile(r"^(\d+/\d+[A-Za-z]*)\.\s+(.+)$")          # "77/52. Subject"
TITLE_HRC_NUM_RE = re.compile(r"^(\d+/\d+[A-Za-z]*)\s+([A-Z].+)$")       # "15/7 Subject"
TITLE_SC_RE = re.compile(r"^Resolution\s+\d+\s*\(\d{4}\)\s*$", re.I)     # "Resolution 1881 (2009)"
TITLE_PRST_RE = re.compile(r"^Statement by the President of the ", re.I)
TITLE_DEC_RE = re.compile(r"^Decision\s+\d+", re.I)

# Operative / subparagraph prefixes (on cleaned, tab-collapsed text).
OP_NUM_RE = re.compile(r"^(\d+)\.\s+(\S.*)$")            # "1. Requests ..."
# Tolerant operative number for source/OCR defects where the period after the
# number is missing or misplaced ("4\tCalls", "232\t.Recognizes", "13\tUrges").
# Only applied mid-operative and only when the clause opens with a known
# operative lead verb, so plain numbers / years never match (see step 11).
OP_NUM_LOOSE_RE = re.compile(r"^(\d{1,3})[.\s]*(\S.*)$")
# Letter cap is 7 to admit long Roman subparagraph markers ("(xxviii)"=6,
# "(xxxviii)"=7). classify_paren() still disambiguates alpha vs Roman, so the
# wider cap only rescues genuine long numerals that {1,5} silently dropped.
OP_PAREN_RE = re.compile(r"^\(([A-Za-z]{1,7}|\d{1,3})\)\s+(\S.*)$")  # "(a) ...", "(i) ...", "(1) ..."

# Frontmatter / structural line patterns.
DIVIDER_RE = re.compile(r"^_{3,}$")
WP_FOOTNOTE_RE = re.compile(r"^(\d{1,3})/\s+(\S.*)$")    # "1/ United Nations, Treaty Series ..."
# Annex/appendix heading. Matches the bare form ("Annex", "Annex II", "Annex A")
# AND a titled form where a label is followed by a running title on the SAME line
# ("Annex A - Items subject to a no-objection process", "ANNEX 1 - <plan>"). A
# title is only accepted AFTER a label (letter / Roman / number) so body prose
# beginning with the word "Annex"/"Appendix" (e.g. "Annex to the present
# resolution ...") is NOT swallowed; the (?=\s|$) also rejects "Annexation ...".
ANNEX_RE = re.compile(
    r"^(Annex|Appendix)(?=\s|$)"
    r"(?:\s+(?:[IVXLCDM]+|[A-Z]|\d{1,3})"
    r"(?:\s*[-–—:.]\s*\S.*|\s+\S.*)?)?\s*$", re.I)
VOTE_RE = re.compile(r"^\[?\s*Adopted\b.*(vote|without a vote)", re.I)
VOTE_TALLY_RE = re.compile(r"^(In favour|Against|Abstaining|Non-Voting|Absent)\s*:\s*(.*)$", re.I)
VOTE_SUMMARY_RE = re.compile(
    r"recorded vote of\s+(\d+)\s+to\s+(\d+)(?:,\s*with\s+(\d+)\s+abstention)?", re.I)
VOTE_KEY = {
    "in favour": "in_favour", "against": "against", "abstaining": "abstaining",
    "non-voting": "non_voting", "absent": "absent",
}
MEETING_RE = re.compile(r"^\d+\s*(st|nd|rd|th)\s+(plenary\s+)?meeting\b", re.I)
# sem-v5 DECISION FORMULA. A GA/ECOSOC *decision* has no "The General Assembly,"
# opening formula: it prints "At its 87th plenary meeting, on 30 June 2023, the
# General Assembly … decided …". Without this the state machine never left `front`
# and the whole decision body was stamped type='frontmatter' (hidden by the site) --
# the mechanism behind the volume-split children rendering as blank pages.
DECISION_OPENING_RE = re.compile(
    r"At\s+(?:its|the)\s+\d{1,4}\s*(?:st|nd|rd|th|d)\s*"
    r"(?:\([^)]{0,24}\)\s*)?"                       # '48th (resumed) plenary meeting'
    r"(?:plenary\s+|resumed\s+|informal\s+|formal\s+)?meeting\b", re.I)
DATE_LINE_RE = re.compile(
    r"^\d{1,2}\s+(January|February|March|April|May|June|July|August|September|"
    r"October|November|December)\s+\d{4}\s*$"
)
SESSION_RE = re.compile(r"session\s*$", re.I)
AGENDA_RE = re.compile(r"^Agenda item\b", re.I)
RUNNING_HEADER_RE = re.compile(r"^[A-Z]+(/[A-Z0-9()./-]+)+$")   # bare doc symbol like "A/RES/48/70"
PAGE_NUM_RE = re.compile(r"^-?\s*\d{1,4}\s*-?$")
# Masthead lines (document letterhead): organisation name, distribution class,
# language, running symbol/date. Only tagged in the front region -- these are
# accounted as frontmatter with subtype='masthead' for a cleaner elements stream.
MASTHEAD_RE = re.compile(
    r"^(United Nations|Nations Unies|UNITED|NATIONS|Distr\.?|GENERAL|LIMITED|"
    r"Original\s*:|ORIGINAL\s*:|Security Council|General Assembly|"
    r"Economic and Social Council|Human Rights Council|Trusteeship Council)\b", re.I)
ADOPTED_BY_RE = re.compile(r"^Adopted by the (Security Council|General Assembly)", re.I)
RES_ADOPTED_RE = re.compile(r"^Resolution adopted by ", re.I)
REPORT_NOTE_RE = re.compile(r"^\[on the (report|recommendation) ", re.I)

# Heading styles seen in the corpus (native docx / doc). TitleH1 / TitleHCH are
# handled as titles; the rest mark section headings inside the body.
BODY_HEADING_STYLES = {"H1", "H2", "H3", "H4", "H23", "H1G", "H2G", "H3G", "H4G",
                       "HCh", "HChG", "HChM"}

# sem-v3 STYLE-TRUSTED HEADINGS: raw style_ids that the document author marked as
# body headings are trusted directly (regardless of bold), mirroring the TOC
# verifier's _HEADING_STYLE list plus the UN G/M-suffixed variants. Matching these
# recovers the section titles the older bold-only heuristic demoted to body text.
STYLE_HEADING_RE = re.compile(
    r"^(H\d+[GM]?|HCh[GM]?|Heading\d+|ArtHead|SectionHead|ChapterHead)$", re.I)
# Numbered resolution-title style -> title tier (handled by _match_title, never a
# body heading). Kept out of STYLE_HEADING_RE deliberately.
STYLE_TITLE_RE = re.compile(r"^(TitleH1|Title\d*)$", re.I)
# Title-ish front-matter styles that must NOT masquerade as body headings.
STYLE_FRONTMATTER_RE = re.compile(r"^(TitleHC[hH]|AgendaTitle.*|Session.*|Distr.*)$", re.I)
# Front-matter TEXT patterns (the verifier's curated FALSE-POSITIVE EXCLUSIONS):
# bracketed committee references, agenda/distribution lines, "Resolution adopted
# by ...", the SC "Adopted by the Security Council at its ..." lines, bare doc
# symbols, and spelled-out session ordinals. A heading-styled line whose text
# matches these stays frontmatter/title, exactly as the verifier excludes it.
STYLE_HEADING_FP_RE = re.compile(
    r"^(\[.*\]$|agenda item\b|distr\b|resolution adopted by\b|adopted by the\b|"
    r"[A-Z]/(RES|PRST|DEC)/|"
    r"(seventy|sixty|fifty|forty|thirty|twenty|nineteen|eighteen|seventeen|sixteen|"
    r"fifteen|fourteen|thirteen|twelfth|eleventh|tenth|ninth|eighth|seventh|sixth|"
    r"fifth|fourth|third|second|first)-?)",
    re.I,
)

# Bold run-in STRUCTURAL LABEL: a whole-paragraph-bold body line that opens with a
# structural label ("Action 13.", "Goal 1.", "Objective 5:") — including the Pact
# for the Future's "Action N. We will ..." commitments (bold label leading a
# sentence). Captured as a heading with the label as prefix.
BOLD_RUNIN_LABEL_RE = re.compile(r"^(action|goal|objective|priority|target)\s+\d+[.:]?\s+\S", re.I)

# Heading prefix extraction (goal 4): a leading section marker split into `prefix`
# so the UI can style it. Labeled ("Objective 1."), delimited marker ("II.", "B.",
# "3)"), or a bare roman numeral leading a title ("I A new generation ...").
_HPREFIX_LABELED = re.compile(
    r"^((?:action|goal|objective|priority|target|article|annex|appendix|chapter|"
    r"section|part|principle|pillar|phase|step)\s+[\dIVXLCM]+[.:]?)\s+(\S.*)$", re.I)
_HPREFIX_MARKER = re.compile(r"^((?:[IVXLCM]{1,6}|[A-Z]|\d{1,3})[.:)])\s+(\S.*)$")
_HPREFIX_ROMAN_NODELIM = re.compile(r"^([IVX]{1,6})\s+([A-Z]\S.*)$")

# A heading whose ENTIRE visible text is just an enumerator marker (roman/letter/
# number), used by the bare-heading + title merge (goal 2).
_BARE_MARKER_RE = re.compile(r"^\(?([IVXLCM]{1,7}|[A-Z]|\d{1,3})\)?\.?$")

# sem-v4 DOTTED ENUMERATOR (SDG-target family): a paragraph whose text OPENS with a
# dotted-numeric ('1.1', '8.10') or letter-suffixed ('1.a') enumerator followed by
# prose. These carry a real sub-section marker INSIDE the text with prefix/level
# NULL (SDG targets under 'Goal N.', forest-plan targets, programme-plan subprogramme
# paragraphs '19.1'..). The marker must move to `prefix`, the clause get a level, and
# the text be stripped. Confirmation is SEQUENCE-BASED (see detect_dotted_enum): a
# lone decimal that merely starts a sentence, an assessment-scale row ('0.01 Albania
# ....'), or a statistics table ('5.5 5.4 Latin America') is NEVER promoted.
DOTTED_ENUM_RE = re.compile(r"^(\d{1,3})\.(\d{1,2}|[a-z])\s+(\S.*)$")
_DOT_LEADER_RE = re.compile(r"\.\s*\.\s*\.")   # TOC / assessment dot leaders

# Split an annex/appendix delimiter into (label, numeral, inline-title). Only ever
# called on lines that already matched ANNEX_RE, so the numeral token is trusted.
_ANNEX_SPLIT_RE = re.compile(
    r"^(Annex|Appendix)(?:\s+([IVXLCDM]+|[A-Z]|\d{1,3})\b)?(?:\s*[-–—:.]\s*|\s+)?(\S.*)?$",
    re.I,
)


def _heading_fp(c: str) -> bool:
    """True if a heading-styled/bold line is really front-matter (verifier FP set)."""
    return bool(
        STYLE_HEADING_FP_RE.match(c) or DATE_LINE_RE.match(c)
        or RUNNING_HEADER_RE.match(c) or PAGE_NUM_RE.match(c)
    )


def _style_is_heading(c: str, lr: "LRow", state: str) -> bool:
    """True if the raw paragraph style marks this as a trusted body heading.

    Front region is skipped (titles/masthead live there); front-matter styles and
    the numbered-title style are excluded; the verifier's FP text patterns keep the
    committee-ref / adopted-by / session lines out even when heading-styled."""
    if state == "front":
        return False
    st = lr.style or ""
    if not st or STYLE_FRONTMATTER_RE.match(st) or STYLE_TITLE_RE.match(st):
        return False
    if not STYLE_HEADING_RE.match(st):
        return False
    return bool(c) and not _heading_fp(c)


def _bold_heading(c: str, lr: "LRow", state: str) -> bool:
    """Whole-paragraph-bold structural label or short title line (verifier bold src)."""
    if state == "front" or not lr.props.get("bold") or not c or _heading_fp(c):
        return False
    if BOLD_RUNIN_LABEL_RE.match(c):
        return True
    # short, non-sentence title fragment -- but never a verb-led operative/preambular
    # clause (those keep their paragraph_type; we do not inflate/deflate labeling).
    if len(c) <= 60 and len(c.split()) <= 8 and not c.endswith("."):
        w0 = c.split(" ", 1)[0].strip(",.").lower()
        if w0 in OPERATIVE_LEAD_VERBS or w0 in PREAMBULAR_FIRST_WORDS:
            return False
        return c[:1].isupper() or c[:1].isdigit()
    return False


def _split_heading_prefix(c: str) -> tuple[str | None, str]:
    """Split a leading section marker into (prefix, text). (None, c) if no marker."""
    for rx in (_HPREFIX_LABELED, _HPREFIX_MARKER, _HPREFIX_ROMAN_NODELIM):
        m = rx.match(c)
        if m:
            return m.group(1).strip(), m.group(2).strip()
    return None, c


def _split_annex_heading(c: str) -> tuple[str, str | None, str | None]:
    """(label, numeral, inline_title) for an annex/appendix delimiter line."""
    m = _ANNEX_SPLIT_RE.match(c)
    if not m:
        return ("Annex", None, None)
    label = m.group(1).title()
    numeral = m.group(2)
    title = (m.group(3) or "").strip() or None
    return (label, numeral, title)


ROMAN_CHARS = set("ivxlcdm")

# Annex that is really an annexed governance INSTRUMENT (terms of reference, rules
# of procedure, statute, charter, ...) with a numbered-article structure gets
# SCOPED operative labeling (its numbered paragraphs are the instrument's
# operatives, tracked independently of the parent resolution). Plain annexes
# (programmes of action, agendas, declarations, lists, schedules, tables) keep
# paragraph_type=None -- their "operativeness" is not resolution-mandate operative
# and is deferred (see run report). "Amendments to ..." annexes are diffs, not the
# instrument, so they are excluded.
ANNEX_INSTRUMENT_RE = re.compile(
    r"\b(terms of reference|rules of procedure|statute|constitution|"
    r"charter of|regulations of|mandate of the)\b", re.I)
ANNEX_AMENDMENT_RE = re.compile(r"^\s*Amendments?\b", re.I)


def _clean(text: str) -> str:
    """Normalize a raw paragraph text for classification & output.

    Collapses tabs to single spaces, turns NBSP into regular spaces, squeezes
    runs of whitespace, and strips. The raw leading-tab indentation is dropped
    (it is captured structurally by props/level), but internal structure is
    preserved as spaces.
    """
    if not text:
        return ""
    t = text.replace("\xa0", " ").replace("\t", " ").replace("\n", " ")
    t = re.sub(r"\s+", " ", t)
    return t.strip()


def _is_terminal(text: str) -> bool:
    """True if `text` ends with sentence/clause-final punctuation."""
    return bool(text) and text.rstrip()[-1:] in ".,;:?!”’)"


# ---------------------------------------------------------------------------
# Logical-row model (raw rows, after WP hard-break merging)
# ---------------------------------------------------------------------------


class LRow:
    """One logical row: a raw row, or several WP rows merged into one clause."""

    __slots__ = ("positions", "kind", "text", "clean", "style", "num", "props",
                 "table_cell", "hyperlinks", "note_ids", "_fr")

    def __init__(self, raw: dict):
        self.positions = [raw["position"]]
        self.kind = raw["kind"]
        self.text = raw["text"] or ""
        self.clean = _clean(self.text)
        self.style = raw["style_id"] or ""
        self.num = raw["numbering"]
        self.props = raw["props"] or {}
        self.table_cell = raw["table_cell"]
        self.hyperlinks = list(raw["hyperlinks"] or [])
        fr = raw["footnote_ref"]
        self.note_ids: list[int] = []
        if isinstance(fr, dict) and "note_ids" in fr:
            self.note_ids = list(fr["note_ids"])
        self._fr = fr

    def merge(self, other: "LRow") -> None:
        self.positions.extend(other.positions)
        joiner = "" if self.text.endswith("-") else " "
        self.text = (self.text.rstrip() + joiner + other.text.lstrip()).strip()
        self.clean = (self.clean + " " + other.clean).strip()
        self.hyperlinks.extend(other.hyperlinks)
        self.note_ids.extend(other.note_ids)


def _structural_start(lr: LRow) -> bool:
    """True if a WP paragraph clearly begins a new structural unit (never merge into prev)."""
    c = lr.clean
    if not c:
        return True
    if OPENING_RE.match(c) or OP_NUM_RE.match(c) or OP_PAREN_RE.match(c):
        return True
    if WP_FOOTNOTE_RE.match(c) or DIVIDER_RE.match(c) or ANNEX_RE.match(c):
        return True
    if VOTE_RE.match(c) or VOTE_TALLY_RE.match(c) or MEETING_RE.match(c):
        return True
    if TITLE_GA_NUM_RE.match(c) or TITLE_SC_RE.match(c) or TITLE_PRST_RE.match(c):
        return True
    if lr.props.get("all_caps") or lr.props.get("alignment") == "center":
        return True
    first = c.split(" ", 1)[0].lower().rstrip(",")
    if first in PREAMBULAR_FIRST_WORDS:
        return True
    return False


def build_logical_rows(raw_rows: list[dict], fmt: str) -> list[LRow]:
    """Merge WP hard-broken continuation lines into single logical rows.

    Conservative: only for wpd, only merges a *body* paragraph into the previous
    body paragraph when the previous line did not end with terminal punctuation
    and the current line is not itself a structural start and looks like a wrap
    (starts lowercase, or lost the firstline indent that the previous line had).
    """
    lrows = [LRow(r) for r in raw_rows]
    if fmt != "wpd":
        return lrows

    merged: list[LRow] = []
    for lr in lrows:
        if (
            merged
            and lr.kind == "paragraph"
            and lr.clean
            and merged[-1].kind == "paragraph"
            and merged[-1].clean
            and not _is_terminal(merged[-1].clean)
            and not _structural_start(lr)
        ):
            prev = merged[-1]
            starts_lower = lr.clean[:1].islower()
            prev_had_indent = "indent_firstline" in prev.props
            cur_no_indent = "indent_firstline" not in lr.props
            if starts_lower or (prev_had_indent and cur_no_indent):
                prev.merge(lr)
                continue
        merged.append(lr)
    return merged


SUBRES_LETTER_RE = re.compile(r"^([A-Z])\.?$")   # bare "A" / "B." sub-resolution heading


def _annex_subtype(lrows: list[LRow], i: int, heading: str) -> str | None:
    """Classify the annex whose heading is lrows[i]: 'amendment', 'instrument', or None.

    Scans from the heading to the next annex/appendix boundary (or end) to read the
    annex's title line and count its numbered paragraphs. Priority:
      * 'amendment' -- the heading OR its title line opens with 'Amendment(s)'
        (e.g. 'Amendments to the terms of reference ...'); the body is a diff
        ('Amend paragraph N to read: ...'), NOT the instrument itself, so it is
        NEVER scoped -- pure labeling, paragraph_type stays null.
      * 'instrument' -- the annex carries its OWN opening formula (an annexed
        resolution/agreement with a preamble), OR its title matches an instrument
        keyword (terms of reference / rules of procedure / statute / ...) and it
        has >=2 numbered paragraphs. Only 'instrument' annexes are scoped (their
        numbered paragraphs are labeled operative, tracked independently).
      * None -- plain annex (programme of action, declaration, agenda, list, ...).
    """
    n = len(lrows)
    title_line = None
    n_numbered = 0
    has_opening = False
    j = i + 1
    while j < n:
        cj = lrows[j].clean
        if lrows[j].kind in ("empty", "section_break") or not cj:
            j += 1
            continue
        if ANNEX_RE.match(cj):           # next annex/appendix -> end of this scope
            break
        if OPENING_RE.match(cj):
            has_opening = True
        if title_line is None:
            title_line = cj
        if OP_NUM_RE.match(cj):
            n_numbered += 1
        j += 1
    if ANNEX_AMENDMENT_RE.match(heading) or (title_line and ANNEX_AMENDMENT_RE.match(title_line)):
        return "amendment"
    if has_opening:
        return "instrument"
    if title_line and ANNEX_INSTRUMENT_RE.search(title_line) and n_numbered >= 2:
        return "instrument"
    return None


def detect_subres_blocks(lrows: list[LRow]) -> dict[int, dict]:
    """Locate consolidated/omnibus sub-resolution boundaries (multi-text).

    A consolidated resolution (e.g. A/RES/48/75 with sub-resolutions A..L) prints,
    for each sub-resolution, a bare capital-letter heading ('A'), then optional
    title line(s), then its OWN opening formula ('The General Assembly,'). We
    segment ONLY when the letter heading is CONFIRMED by an opening formula that
    follows within a short window -- so section headings inside a single resolution
    ('I', 'A. Utilization ...' followed by '1. ...', never an opening) are never
    mistaken for a new text.

    Returns {lrow_index_of_letter: {'letter': str, 'title_idx': [lrow_index, ...]}}.
    The title lines are the non-empty rows between the letter and the opening; they
    are merged into ONE title element per block by the caller and skipped in the
    normal stream (their positions are consumed by that title element).
    """
    n = len(lrows)
    blocks: dict[int, dict] = {}
    for i, lr in enumerate(lrows):
        m = SUBRES_LETTER_RE.match(lr.clean)
        if not m:
            continue
        title_idx: list[int] = []
        seen_open = False
        j = i + 1
        steps = 0
        while j < n and steps < 6:
            cj = lrows[j].clean
            if lrows[j].kind in ("empty", "section_break") or not cj:
                j += 1
                continue
            if OPENING_RE.match(cj):
                seen_open = True
                break
            # anything that is clearly not a sub-title stops the look-ahead
            if (OP_NUM_RE.match(cj) or OP_PAREN_RE.match(cj) or ANNEX_RE.match(cj)
                    or MEETING_RE.match(cj) or VOTE_RE.match(cj) or DIVIDER_RE.match(cj)
                    or WP_FOOTNOTE_RE.match(cj) or SUBRES_LETTER_RE.match(cj)
                    or TITLE_GA_NUM_RE.match(cj)):
                break
            title_idx.append(j)
            j += 1
            steps += 1
        if seen_open:
            blocks[i] = {"letter": m.group(1), "title_idx": title_idx}
    return blocks


def _dotted_enum_candidate(c: str) -> tuple[int, str, int, str, str] | None:
    """(major, kind, ord, prefix, rest) if `c` opens with a plausible dotted
    enumerator ('1.1', '8.10', '1.a') followed by PROSE, else None.

    Pure lexical screen (no sequence context yet): the regex self-rejects tabular
    decimals ('2.983 745,750 ...' — no space after a 2+digit minor), and the guards
    drop major-0 assessment rows ('0.01 Albania'), dot-leader/statistics rows, and
    numeric-heavy fragments. detect_dotted_enum() then confirms only those that sit
    in a coherent contiguous run, so a stray decimal opening a sentence never wins.
    """
    m = DOTTED_ENUM_RE.match(c)
    if not m:
        return None
    major = int(m.group(1))
    if major < 1:                      # assessment scales are 0.xx — never enumerators
        return None
    minor = m.group(2)
    rest = m.group(3).strip()
    if len(rest) < 15 or not rest[:1].isalpha():
        return None
    if _DOT_LEADER_RE.search(rest):    # 'Albania ....... .' country/TOC leaders
        return None
    head = rest[:40]
    if sum(ch.isalpha() for ch in head) / len(head) < 0.55:   # numeric-heavy table row
        return None
    prefix = f"{major}.{minor}"
    if minor.isdigit():
        return (major, "num", int(minor), prefix, rest)
    return (major, "let", ord(minor) - ord("a") + 1, prefix, rest)


def _dotted_follows(a: tuple, b: tuple) -> bool:
    """True if candidate b directly CONTINUES candidate a's enumerator sequence.

    Same major: num->num (b=a+1), num->letter 'a' (1.5 -> 1.a), letter->letter next.
    New major: only major+1 restarting at numeric .1 (1.b -> 2.1)."""
    (amaj, akind, aord, *_), (bmaj, bkind, bord, *_) = a, b
    if bmaj == amaj:
        if akind == "num" and bkind == "num":
            return bord == aord + 1
        if akind == "num" and bkind == "let":
            return bord == 1
        if akind == "let" and bkind == "let":
            return bord == aord + 1
        return False
    if bmaj == amaj + 1:
        return bkind == "num" and bord == 1
    return False


def detect_dotted_enum(lrows: list["LRow"]) -> dict[int, tuple[str, int, str]]:
    """Map lrow-index -> (prefix, level, stripped_text) for CONFIRMED dotted
    enumerators (sem-v4). Confirmation = sequence: the candidate belongs to a
    contiguous follows()-chain of length >=3 that STARTS at a numeric '.1'. This is
    the SDG-target / forest-goal / programme-plan family; assessment tables and lone
    decimals fail to chain and are left untouched (prefix/level stay NULL).
    """
    cand: list[tuple[int, tuple]] = []
    for i, lr in enumerate(lrows):
        if lr.kind != "paragraph" or not lr.clean:
            continue
        c = _dotted_enum_candidate(lr.clean)
        if c is not None:
            cand.append((i, c))
    confirmed: dict[int, tuple[str, int, str]] = {}
    k, m = 0, len(cand)
    while k < m:
        chain = [cand[k]]
        j = k + 1
        while j < m and _dotted_follows(cand[j - 1][1], cand[j][1]):
            chain.append(cand[j])
            j += 1
        # a genuine enumerator run is >=3 long and opens at a numeric '.1'
        head_kind, head_ord = chain[0][1][1], chain[0][1][2]
        if len(chain) >= 3 and head_kind == "num" and head_ord == 1:
            for idx, cc in chain:
                _, _, _, prefix, rest = cc
                confirmed[idx] = (prefix, 2, rest)
        k = j
    return confirmed


# ---------------------------------------------------------------------------
# sem-v5: MARKER SEQUENCE CONFIRMATION
#
# The raw layer is not trustworthy about leading numbers. The PDF extractor merges
# hanging numbers back onto the text that follows them, and a page number, a
# footnote reference or a column-margin artifact merges exactly like a real
# operative marker does; the parser then stored it as an operative number that is
# not in the source (263 such fabrications corpus-wide; `S/RES/661(1990)` held a
# '19.' on a mid-sentence continuation of operative 3(c)). A marker is therefore
# only accepted when the document's OWN numbering vouches for it — the same
# discipline detect_dotted_enum() applies to SDG-style enumerators:
#
#   * it RESTARTS a list ('1.'), or
#   * it CONTINUES a run already seen (1..MARKER_MAX_GAP ahead of a live value —
#     the gap tolerates numbers the source itself dropped, cf.
#     RESCUE_INFERRED_OPERATIVE), or
#   * it OPENS a run that the next few candidates confirm (a '+1' successor within
#     MARKER_LOOKAHEAD candidates), or
#   * it reads unmistakably like an operative clause (capitalised finite operative
#     lead verb) AND lies inside the band the document's own operative count allows.
#
# Nothing else is a marker. A refused number is NOT deleted: it stays inside the
# element's text (a gap in labelling is recoverable, an invented legal citation is
# a false statement), the element gets no prefix and no operative labeling, and the
# refusal is recorded in `issues[]` so it is queryable in SQL after the fact.
# ---------------------------------------------------------------------------

MARKER_MAX_GAP = 3        # forward step tolerated inside a live run (source-dropped numbers)
MARKER_LOOKAHEAD = 3      # candidates scanned ahead for a '+1' confirmation
MARKER_BAND_SLACK = 3     # a marker may exceed the document's own operative count by <= 3

# The remainder after a candidate marker that proves it is NOT an operative clause:
# an old-style resolution heading printed in a volume / table of contents
# ("1110 (XI). Admission of Morocco…", "(XI). …") — the leading number is that
# entry's PAGE number, not a paragraph marker.
_OLD_STYLE_HEADING_RE = re.compile(r"^\d{0,4}\s*\([IVXLCDM]+\)\s*\.")
# An upper-case Roman parenthetical: never a subparagraph marker in this corpus
# unless the sequence confirms it ('(IX) of 11 December 1954' is the tail of
# 'resolution 900 (IX) of 11 December 1954' broken across lines).
_UPPER_ROMAN_RE = re.compile(r"^[IVXLCDM]{1,7}$")

_ROMAN_VALUES = {"i": 1, "v": 5, "x": 10, "l": 50, "c": 100, "d": 500, "m": 1000}


def _roman_value(tok: str) -> int:
    """Value of a Roman numeral token ('xiv' -> 14). 0 if not a Roman numeral."""
    t = tok.lower()
    if not t or any(ch not in _ROMAN_VALUES for ch in t):
        return 0
    total = prev = 0
    for ch in reversed(t):
        v = _ROMAN_VALUES[ch]
        total += -v if v < prev else v
        prev = max(prev, v)
    return total


def _op_verb_start(text: str) -> bool:
    """True if `text` opens with a capitalised finite operative lead verb.

    'Decides that ...', 'Also requests ...' -> True; 'nationals or in their
    territories ...' (a mid-sentence continuation) -> False."""
    t = text.strip()
    if not t[:1].isupper():
        return False
    words = t.split()
    w0 = words[0].strip(",.").lower()
    if w0 in ("also", "further", "again", "finally", "moreover") and len(words) > 1:
        w0 = words[1].strip(",.").lower()
    return w0 in OPERATIVE_LEAD_VERBS


def _num_marker_candidate(lr: "LRow") -> tuple[int, str] | None:
    """(number, remaining text) if this row opens with a plausible 'N.' marker.

    Pure lexical screen, no sequence context yet. Rejects the numbered resolution
    TITLE form ('77/52. Subject'), volume/TOC heading lines whose leading number is
    a page number ('1110 (XI). Admission of Morocco…'), and dot-leader TOC rows.
    """
    if lr.kind != "paragraph":
        return None
    c = lr.clean
    if not c or TITLE_GA_NUM_RE.match(c):
        return None
    m = OP_NUM_RE.match(c)
    if not m:
        return None
    rest = m.group(2).strip()
    if _OLD_STYLE_HEADING_RE.match(rest) or _DOT_LEADER_RE.search(rest):
        return None
    return int(m.group(1)), rest


def _marker_band(lrows: list["LRow"]) -> int:
    """Highest number this document could legitimately have printed as a marker.

    Two bounds, whichever is larger, both taken from the parser's INPUT (the raw
    rows) and never from what the parse produced, so a document can never vouch for
    a number by having stored it:
      * its operative-shaped rows -- numeric marker candidates plus rows opening
        with a finite operative verb;
      * its body rows -- a document cannot number more paragraphs than it has rows.
    Plus MARKER_BAND_SLACK for numbers the source itself dropped. What this rejects
    is the page-number family: a table-of-contents page of 35 rows cannot carry
    markers '59.'-'64.' (`A/RES/1110(XI)`), and a 12-row decision cannot carry
    '177.' (`E/DEC/2010/235`).
    """
    op_shaped = body = 0
    for lr in lrows:
        if lr.kind != "paragraph" or not lr.clean:
            continue
        body += 1
        if _num_marker_candidate(lr) is not None or _op_verb_start(lr.clean):
            op_shaped += 1
    return max(op_shaped, body) + MARKER_BAND_SLACK


def detect_op_numbers(lrows: list["LRow"]) -> tuple[dict[int, int], list[tuple[int, int, str]]]:
    """Confirm which leading 'N.' numbers are real operative markers.

    Returns ({lrow_index: number} for CONFIRMED markers, [(lrow_index, number,
    text_head)] for REFUSED ones). See the section header for the rule. Live runs
    are cleared at an opening formula or an annex delimiter, so a second text block
    never inherits the first block's numbering.
    """
    cands: list[tuple[int, int, str]] = []      # (lrow index, number, rest)
    resets: set[int] = set()                    # candidate ordinals that start fresh
    for i, lr in enumerate(lrows):
        c = lr.clean
        if not c:
            continue
        if OPENING_RE.match(c) or ANNEX_RE.match(c):
            resets.add(len(cands))
            continue
        cand = _num_marker_candidate(lr)
        if cand is not None:
            cands.append((i, cand[0], cand[1]))

    band = _marker_band(lrows)
    confirmed: dict[int, int] = {}
    refused: list[tuple[int, int, str]] = []
    live: dict[int, int] = {}                   # last value seen -> candidate ordinal
    for k, (idx, num, rest) in enumerate(cands):
        if k in resets:
            live.clear()
        ok = False
        if num > band:
            ok = False                          # outside the document's own numbering
        elif num == 1:
            ok = True                           # a list may always restart
        else:
            for gap in range(1, MARKER_MAX_GAP + 1):
                if (num - gap) in live:
                    live.pop(num - gap)
                    ok = True
                    break
            if not ok:
                for j in range(k + 1, min(len(cands), k + 1 + MARKER_LOOKAHEAD)):
                    if cands[j][1] == num + 1:
                        ok = True
                        break
            if not ok and _op_verb_start(rest):
                ok = True                       # unmistakable operative clause
        if ok:
            live[num] = k
            confirmed[idx] = num
        else:
            refused.append((idx, num, rest[:60]))
    return confirmed, refused


def detect_upper_roman_parens(lrows: list["LRow"]) -> set[int]:
    """Indices of '(IX)'-style UPPER-CASE Roman parentheticals that are confirmed
    subparagraph markers (they continue, or open a run the next candidates confirm).

    Everything else with that shape is the tail of a citation whose line broke after
    the session number ('… resolution 900 (IX) of 11 December 1954'), which used to
    be stored as a subparagraph marker in 815 rows across 677 documents."""
    cands: list[tuple[int, int]] = []
    for i, lr in enumerate(lrows):
        if lr.kind != "paragraph" or not lr.clean:
            continue
        m = OP_PAREN_RE.match(lr.clean)
        if m and _UPPER_ROMAN_RE.match(m.group(1)):
            cands.append((i, _roman_value(m.group(1))))
    confirmed: set[int] = set()
    last: int | None = None
    for k, (idx, val) in enumerate(cands):
        ok = last is not None and val == last + 1
        if not ok:
            for j in range(k + 1, min(len(cands), k + 1 + MARKER_LOOKAHEAD)):
                if cands[j][1] == val + 1:
                    ok = True
                    break
        if ok:
            confirmed.add(idx)
            last = val
    return confirmed


# ---------------------------------------------------------------------------
# Sub-paragraph level / (i) disambiguation
# ---------------------------------------------------------------------------


class OpLevelTracker:
    """Assigns level/prefix to operative markers, disambiguating (i) alpha vs roman."""

    def __init__(self) -> None:
        self.reset()

    def reset(self) -> None:
        self.last_alpha: str | None = None   # last single-letter alpha subpara
        self.roman_active = False            # a roman sub-sub sequence is running

    def top(self) -> None:
        # New top-level operative resets sub-sequences.
        self.last_alpha = None
        self.roman_active = False

    def classify_paren(self, token: str, next_tok: str | None = None) -> tuple[int, str]:
        """Return (level, prefix) for a parenthetical marker token like 'a','i','iv','1'.

        `next_tok` is the following paren marker (if any); it disambiguates a single
        '(i)' -- Roman nesting when '(ii)' follows, alpha continuation otherwise.
        """
        tok = token.lower()
        prefix = f"({token})"
        if tok.isdigit():
            return 2, prefix  # numeric subpara -- treat as level 2 variant
        # DOUBLED same-letter marker ('aa','bb',..,'zz'): the UN alpha convention
        # for subparagraphs continuing past (z). These are ALWAYS level-2 alpha --
        # NOT Roman -- even when the letter happens to be a Roman glyph ('cc','dd',
        # 'ii','ll','mm' etc.). The only exception is a genuine Roman sub-sub run
        # already open (e.g. '(i)(ii)') where '(ii)' is Roman: gated on roman_active.
        if len(tok) == 2 and tok[0] == tok[1]:
            if self.roman_active and all(ch in ROMAN_CHARS for ch in tok):
                return 3, prefix
            self.last_alpha = tok
            self.roman_active = False
            return 2, prefix
        # multi-char: roman if all roman chars, else treat as deeper alpha
        if len(tok) > 1:
            if all(ch in ROMAN_CHARS for ch in tok):
                self.roman_active = True
                return 3, prefix
            return 2, prefix
        # single char
        if tok == "i":
            # '(i)' is ambiguous: the start of a Roman sub-sub run (i)(ii)(iii)…
            # nested under an alpha item, OR alpha continuation of a flat (a)(b)…(i)
            # list (incl. …(d),(i) when source dropped (e)-(h)). Disambiguate by the
            # NEXT marker: '(ii)' => Roman nesting; otherwise, if an alpha run is
            # open, alpha continuation.
            if next_tok is not None and next_tok.lower() == "ii":
                self.roman_active = True
                return 3, prefix
            if self.last_alpha is not None and not self.roman_active:
                self.last_alpha = "i"
                return 2, prefix
            self.roman_active = True         # start of (i)(ii)... roman run
            return 3, prefix
        if tok in ("v", "x") and self.roman_active:
            return 3, prefix
        if tok in ROMAN_CHARS and self.roman_active and self.last_alpha is None:
            return 3, prefix
        # default: alpha subparagraph
        self.last_alpha = tok
        self.roman_active = False
        return 2, prefix


# ---------------------------------------------------------------------------
# Element construction
# ---------------------------------------------------------------------------


def _new_element(lr: LRow, **kw) -> dict:
    el = {
        "positions": list(lr.positions),
        "type": kw.get("type", "paragraph"),
        "section": kw.get("section", "main"),
        "paragraph_type": kw.get("paragraph_type"),
        "level": kw.get("level"),
        "prefix": kw.get("prefix"),
        "heading_level": kw.get("heading_level"),
        "text": kw.get("text", lr.clean),
        "lead_verb": kw.get("lead_verb"),
        "hyperlinks": lr.hyperlinks or [],
        "note_ids": sorted(set(lr.note_ids)) if lr.note_ids else [],
    }
    if kw.get("subtype"):
        el["subtype"] = kw["subtype"]
    if kw.get("annex_index"):
        el["annex_index"] = kw["annex_index"]
    if kw.get("text_index", 1) != 1:
        el["text_index"] = kw["text_index"]
    return el


def _lead_verb_from_text(clean: str) -> str | None:
    """Lexical lead-verb extraction for WP (no italics). Returns the verb phrase.

    Keeps the phrase tight: the leading participle plus a meaningful modifier
    only (e.g. "Recalling also", "Deeply concerned", "Bearing in mind", "Taking
    note"), not trailing objects -- "Recalling that ..." yields just "Recalling".
    """
    words = [w.strip(",") for w in clean.split()]
    if not words:
        return None
    w0 = words[0].lower()
    if w0 not in PREAMBULAR_FIRST_WORDS:
        return None
    if w0 in ("deeply", "gravely", "fully", "firmly", "strongly", "keenly") and len(words) > 1:
        return f"{words[0]} {words[1]}"          # "Deeply concerned"
    if w0 == "bearing" and len(words) > 2:
        return "Bearing in mind"
    if w0 == "taking" and len(words) > 1:
        return f"Taking {words[1]}"              # "Taking note"
    if w0 == "having" and len(words) > 1:
        return f"Having {words[1]}"              # "Having considered"
    if w0 == "guided":
        return "Guided by"
    if len(words) > 1 and words[1].lower() in ("also", "further", "again"):
        return f"{words[0]} {words[1]}"          # "Recalling also"
    return words[0]


def _split_countries(s: str) -> list[str]:
    """Split a comma-separated country list, dropping trailing bracket/period."""
    s = s.strip().rstrip("]").rstrip(".").strip()
    if not s:
        return []
    return [p.strip() for p in s.split(",") if p.strip()]


def _consume_vote_block(lrows: list["LRow"], i: int) -> tuple[int, list[int], str, dict, dict | None]:
    """Greedily consume a vote block starting at lrows[i].

    Returns (positions, summary_text, vote, vote_summary). The block is the
    optional "[Adopted by a recorded vote ...]" line plus the
    "In favour:/Against:/Abstaining:" labels and their country lists (whether
    inline or split across following rows). Consumption stops at the first
    blank/footnote/divider/meeting/opening/annex/operative/date row, or at a
    non-label row while no tally label is currently open (so we never swallow
    unrelated tail prose)."""
    n = len(lrows)
    positions: list[int] = []
    summary_parts: list[str] = []
    vote: dict[str, list[str]] = {"in_favour": [], "against": [], "abstaining": []}
    vote_summary: dict | None = None
    cur: str | None = None
    j = i
    started = False
    while j < n:
        r = lrows[j]
        rc = r.clean
        if r.kind in ("empty", "section_break", "footnote") or not rc:
            break
        if (DIVIDER_RE.match(rc) or MEETING_RE.match(rc) or OPENING_RE.match(rc)
                or ANNEX_RE.match(rc) or OP_NUM_RE.match(rc) or DATE_LINE_RE.match(rc)):
            break
        m_label = VOTE_TALLY_RE.match(rc)
        is_adopted = bool(VOTE_RE.match(rc))
        if started and not (m_label or is_adopted) and cur is None:
            break
        if is_adopted:
            summary_parts.append(rc)
            ms = VOTE_SUMMARY_RE.search(rc)
            if ms:
                vote_summary = {"in_favour": int(ms.group(1)), "against": int(ms.group(2)),
                                "abstaining": int(ms.group(3)) if ms.group(3) else 0}
        elif m_label:
            cur = VOTE_KEY.get(m_label.group(1).lower())
            if cur is not None:
                vote.setdefault(cur, [])
                rest = m_label.group(2).strip()
                if rest:
                    vote[cur].extend(_split_countries(rest))
        elif cur is not None:
            vote[cur].extend(_split_countries(rc))
        positions.extend(r.positions)
        j += 1
        started = True
    summary = " ".join(summary_parts) if summary_parts else "Vote record"
    return j, positions, summary, vote, vote_summary


def _footnote_text(raw_text: str) -> str:
    """Clean a footnote body: drop leading tabs and marker glyphs (*, ?, N/, N.)."""
    t = _clean(raw_text)
    t = re.sub(r"^[\*\?†‡\s]+", "", t)         # symbol markers
    t = re.sub(r"^\d{1,3}[/.]\s*", "", t)                 # "1/ " or "1. "
    t = re.sub(r"^[\*\?†‡]+\s*", "", t)
    return t.strip()


# ---------------------------------------------------------------------------
# Action-verb annotation pass (migration 004 / sem-v2)
# ---------------------------------------------------------------------------


def annotate_actions(elements: list[dict]) -> None:
    """Attach a nested ``action`` dict to each operative/preambular element in place.

    Runs the deterministic action-verb parser (``fulltext_verbs.extract_action``)
    over the element sequence in document order, resolving each clause's *chapeau*
    context EXACTLY as ``fulltext_verbs_eval.run_parser_over_doc`` does (this is the
    single source of truth for that resolution — keep the two in lockstep):

      * only ``paragraph_type IN ('operative','preambular')`` elements are annotated;
        every other element (PRST/statement bodies, annexes, headings, votes, and
        every ``paragraph_type IS NULL`` element) is left untouched — no ``action`` key;
      * a top-level (level<=1) clause that does NOT end in ':' resets the inherited
        chapeau context;
      * the row a sub-item inherits from is the most recent operative element ending
        in ':' at a shallower level — with ``governing_verb_for_children`` overriding
        the line's own leading verb for declaration ('We decide to:') and passive
        personified ('... are encouraged ... :') chapeaux.

    Elements whose ``extract_action`` returns ``None`` (noun-phrase budget sub-items,
    chapeau-less continuations, non-action headings that slipped through) get no
    ``action`` key and so load as all-NULL action columns. The Chapter-VII marker
    ('Acting under Chapter VII …') IS attached (it has a context_marker) even though
    its normalized verb is None. The pass never alters ``text``/``positions`` — it is
    purely additive, so the accounting invariant and text-preservation gate are
    unaffected.
    """
    chapeau: dict | None = None
    for el in elements:
        ptype = el.get("paragraph_type")
        if ptype not in ("operative", "preambular"):
            continue
        text = el.get("text") or ""
        level = el.get("level")
        prefix = el.get("prefix")
        is_colon = text.rstrip().endswith(":")
        lvl = level if level is not None else (1 if ptype == "operative" else 0)
        # a new top-level, non-chapeau clause resets inherited context
        if lvl <= 1 and not is_colon:
            chapeau = None
        action = fv.extract_action(
            text, paragraph_type=ptype, level=level, prefix=prefix,
            chapeau_action=chapeau,
        )
        # this row becomes the chapeau for following sub-items if it opens a list;
        # a governing verb overrides the line's own leading verb for what children
        # inherit ('We decide to:' -> decide; '... are encouraged to ... :' -> encourage).
        gov = fv.governing_verb_for_children(text)
        if gov is not None:
            chapeau = gov
        elif (action and not action.get("inherited") and is_colon
                and action.get("normalized")):
            chapeau = action
        if action is not None:
            el["action"] = action


# ---------------------------------------------------------------------------
# Main per-document parse
# ---------------------------------------------------------------------------


_MERGE_SKIP_SUBTYPES = {"subres", "annex", "appendix", "instrument", "amendment"}


def _is_bare_marker_heading(el: dict) -> bool:
    """A heading whose only content is a section marker (its text, or a marker-only
    prefix with empty text). Annex/subres delimiters are excluded."""
    if el.get("type") != "heading" or el.get("subtype") in _MERGE_SKIP_SUBTYPES:
        return False
    txt = (el.get("text") or "").strip()
    if txt:
        return bool(_BARE_MARKER_RE.match(txt))
    pfx = (el.get("prefix") or "").strip().rstrip(".)")
    return bool(pfx) and bool(_BARE_MARKER_RE.match(pfx))


def _is_title_like(el: dict) -> bool:
    """The following element reads like the title that belongs to a bare marker: a
    non-marker heading, or a short non-sentence NULL-type paragraph."""
    t = (el.get("text") or "").strip()
    if not t or _BARE_MARKER_RE.match(t):
        return False
    if el.get("type") == "heading":
        return True
    if el.get("type") == "paragraph" and el.get("paragraph_type") is None:
        return len(t.split()) <= 14 and t[-1:] not in ".;" and (t[:1].isupper() or t[:1].isdigit())
    return False


def _merge_bare_headings(elements: list[dict]) -> list[dict]:
    """Goal 2: merge a bare section-marker heading ('I', 'II.', 'A.') with the short
    title line that immediately follows it (within the same block) into ONE heading:
    prefix='I.', text='General guidelines'. Positions/hyperlinks/notes are unioned."""
    out: list[dict] = []
    i, n = 0, len(elements)
    while i < n:
        el = elements[i]
        if i + 1 < n and _is_bare_marker_heading(el):
            nxt = elements[i + 1]
            same_block = (
                el.get("section") == nxt.get("section")
                and el.get("text_index", 1) == nxt.get("text_index", 1)
                and el.get("annex_index") == nxt.get("annex_index"))
            if same_block and _is_title_like(nxt):
                marker = ((el.get("text") or "").strip()
                          or (el.get("prefix") or "").strip()).rstrip(".)").strip()
                ntext = ((nxt.get("prefix") or "") + (nxt.get("text") or "")).strip()
                merged = dict(el)
                merged["prefix"] = f"{marker}." if marker else None
                merged["text"] = ntext
                merged["positions"] = list(el.get("positions") or []) + list(nxt.get("positions") or [])
                merged["hyperlinks"] = (el.get("hyperlinks") or []) + (nxt.get("hyperlinks") or [])
                merged["note_ids"] = sorted(set((el.get("note_ids") or []) + (nxt.get("note_ids") or [])))
                out.append(merged)
                i += 2
                continue
        out.append(el)
        i += 1
    return out


def _stamp_annex_index(elements: list[dict]) -> None:
    """Annex-contract: every element in section='annex' carries the annex_index of
    its delimiter (delimiters set it in the loop; body elements inherit it here)."""
    cur: int | None = None
    for el in elements:
        if el.get("section") == "annex":
            if el.get("annex_index") is not None:
                cur = el["annex_index"]
            elif cur is not None:
                el["annex_index"] = cur
        elif el.get("section") == "main":
            cur = None  # a new main-section block (opening / sub-res) ends annex scope


def parse_document(symbol: str, fmt: str, raw_rows: list[dict]) -> dict:
    lrows = build_logical_rows(raw_rows, fmt)
    subres = detect_subres_blocks(lrows)
    dotted_enum = detect_dotted_enum(lrows)  # sem-v4: confirmed N.N / N.a enumerators
    # sem-v5: only numbers the document's own sequence vouches for are markers.
    op_numbers, refused_markers = detect_op_numbers(lrows)
    refused_idx = {r[0] for r in refused_markers}
    upper_roman_parens = detect_upper_roman_parens(lrows)
    # lrow indices whose positions are consumed by a block's merged title element
    title_skip: set[int] = set()
    for blk in subres.values():
        title_skip.update(blk["title_idx"])

    elements: list[dict] = []
    dropped: list[dict] = []
    issues: list[dict] = []

    state = "front"          # front -> preamble -> operative ; plus tail signals
    section = "main"
    annex_index = 0
    text_index = 1
    seen_opening = False
    seen_title = False
    pending_block_opening = False   # a sub-res boundary already bumped text_index;
                                    # suppress the imminent opening's own bump
    annex_scoped = False            # current annex/appendix is a scoped instrument
                                    # (its numbered paras are labeled operative)
    last_op_number = 0              # last top-level operative number seen (for rescue)
    last_heading_level = None        # tier of the most recent section heading, so a
                                     # bold run-in ('Action N.') nests one level below
    op_tracker = OpLevelTracker()

    i = 0
    n = len(lrows)
    while i < n:
        lr = lrows[i]
        c = lr.clean
        kind = lr.kind

        # 0. sub-res block title lines are emitted with their block's letter
        # heading (below); skip them here so positions are consumed exactly once.
        if i in title_skip:
            i += 1
            continue

        # 0b. sub-resolution block boundary (letter heading + title + opening ahead)
        if i in subres:
            blk = subres[i]
            if seen_opening:
                text_index += 1
                pending_block_opening = True
            op_tracker.reset()
            op_tracker.top()
            section = "main"
            annex_index = 0
            state = "blockhead"
            elements.append(_new_element(lr, type="heading", section="main",
                                         heading_level=1, text=c,
                                         text_index=text_index, subtype="subres"))
            tidx = blk["title_idx"]
            if tidx:
                t_positions: list[int] = []
                t_texts: list[str] = []
                t_hyper: list = []
                for k in tidx:
                    t_positions.extend(lrows[k].positions)
                    if lrows[k].clean:
                        t_texts.append(lrows[k].clean)
                    t_hyper.extend(lrows[k].hyperlinks)
                title_el = {
                    "positions": t_positions, "type": "title", "section": "main",
                    "paragraph_type": None, "level": None, "prefix": None,
                    "heading_level": None, "text": " ".join(t_texts),
                    "lead_verb": None, "hyperlinks": t_hyper, "note_ids": [],
                }
                if text_index != 1:
                    title_el["text_index"] = text_index
                elements.append(title_el)
                seen_title = True
            i += 1
            continue

        # 1. empties & section breaks -----------------------------------------
        if kind == "empty" or (kind != "footnote" and not c):
            reason = "empty" if kind in ("empty", "paragraph") else kind
            for p in lr.positions:
                dropped.append({"position": p, "reason": reason})
            i += 1
            continue
        if kind == "section_break":
            for p in lr.positions:
                dropped.append({"position": p, "reason": "section_break"})
            i += 1
            continue

        # 2. native/doc footnote rows -----------------------------------------
        if kind == "footnote":
            note_id = None
            if isinstance(lr._fr, dict):
                note_id = lr._fr.get("note_id")
            el = _new_element(lr, type="footnote", section=section, text=_footnote_text(lr.text),
                              text_index=text_index)
            if note_id is not None:
                el["note_ids"] = [note_id]
            else:
                # sem-v5: _footnote_text strips the leading marker glyph ('7 ', '1/ ',
                # '3. '). Keep it as the note id instead of deleting it -- the
                # citation link survives and the word conservation check balances.
                m_mark = re.match(r"^\s*(\d{1,3})\s*[/.]?\s+\S", lr.clean)
                if m_mark:
                    el["note_ids"] = [int(m_mark.group(1))]
            elements.append(el)
            i += 1
            continue

        # 3. table cells: group consecutive cells -----------------------------
        if kind == "table_cell":
            j = i
            group: list[LRow] = []
            while j < n and lrows[j].kind == "table_cell":
                group.append(lrows[j])
                j += 1
            positions = [p for g in group for p in g.positions]
            cells = [g.clean for g in group if g.clean]
            # masthead (before we've reached the body) vs genuine data table
            is_masthead = (state == "front" and not seen_title
                           and len(cells) <= 12)
            if not cells:
                for p in positions:
                    dropped.append({"position": p, "reason": "layout_cell"})
            else:
                el = {
                    "positions": positions,
                    "type": "frontmatter" if is_masthead else "table",
                    "section": section,
                    "paragraph_type": None,
                    "level": None,
                    "prefix": None,
                    "heading_level": None,
                    "text": " | ".join(cells),
                    "lead_verb": None,
                    "hyperlinks": [h for g in group for h in g.hyperlinks],
                    "note_ids": [],
                }
                if is_masthead:
                    el["subtype"] = "masthead"
                if text_index != 1:
                    el["text_index"] = text_index
                elements.append(el)
            i = j
            continue

        # 4. divider (footnote separator) -------------------------------------
        if DIVIDER_RE.match(c):
            elements.append(_new_element(lr, type="divider", section=section, text=c,
                                         text_index=text_index))
            i += 1
            continue

        # 5. WP inline footnote "N/ text" -------------------------------------
        if fmt == "wpd" and WP_FOOTNOTE_RE.match(c):
            m = WP_FOOTNOTE_RE.match(c)
            el = _new_element(lr, type="footnote", section=section,
                              text=m.group(2).strip(), text_index=text_index)
            el["note_ids"] = [int(m.group(1))]
            elements.append(el)
            i += 1
            continue

        # 6. opening formula --------------------------------------------------
        if OPENING_RE.match(c):
            if seen_opening and not pending_block_opening:
                # repeated opening with NO preceding letter heading (e.g. an ECOSOC
                # resolution that recommends a GA text): still a distinct text block.
                text_index += 1
                op_tracker.reset()
                section = "main"
                annex_index = 0
            pending_block_opening = False
            seen_opening = True
            state = "preamble"
            elements.append(_new_element(lr, type="opening", section=section,
                                         paragraph_type="preambular", level=0,
                                         text=c, text_index=text_index))
            i += 1
            continue

        # 7. annex / appendix heading -> new section --------------------------
        # The delimiter is a heading carrying subtype ('annex'/'appendix', or the
        # scoped 'instrument'/'amendment'), prefix = the label+numeral ('Annex II'),
        # and text = the annex TITLE (inline after a dash, or merged from the next
        # line in step 7b). last_heading_level is reset so run-ins nest below it.
        if ANNEX_RE.match(c) and state != "front":
            m = ANNEX_RE.match(c)
            if m.group(1).lower() == "annex":
                annex_index += 1
                section = "annex"
            else:
                section = "appendix"
            op_tracker.top()
            op_tracker.reset()
            label, numeral, inline_title = _split_annex_heading(c)
            annex_prefix = (label + (f" {numeral}" if numeral else "")).strip()
            annex_sub = _annex_subtype(lrows, i, c)
            annex_scoped = annex_sub == "instrument"
            subtype = annex_sub or ("appendix" if section == "appendix" else "annex")
            el = _new_element(lr, type="heading", section=section, heading_level=1,
                              prefix=annex_prefix, text=inline_title or "",
                              subtype=subtype, text_index=text_index)
            if section == "annex":
                el["annex_index"] = annex_index
            elements.append(el)
            last_heading_level = 1
            # 'annextitle' folds the following title line into this delimiter (below);
            # then the (scoped) preamble/operative machine runs.
            state = "annextitle"
            i += 1
            continue

        # 7b. annex title line: merge the first plain title line INTO the delimiter
        # heading (goal / annex-contract: one heading element carries label+title).
        if state == "annextitle":
            structural = bool(
                OPENING_RE.match(c) or OP_NUM_RE.match(c) or OP_PAREN_RE.match(c)
                or ANNEX_RE.match(c) or MEETING_RE.match(c) or DIVIDER_RE.match(c)
                or VOTE_RE.match(c) or VOTE_TALLY_RE.match(c))
            delim = elements[-1] if elements else None
            can_merge = (
                delim is not None and delim.get("type") == "heading"
                and delim.get("section") in ("annex", "appendix")
                and not (delim.get("text") or "").strip())
            hp, _ = _split_heading_prefix(c)
            if not structural and can_merge and hp is None and len(c) <= 200:
                delim["positions"] = list(delim["positions"]) + list(lr.positions)
                delim["text"] = c
                if lr.hyperlinks:
                    delim["hyperlinks"] = (delim.get("hyperlinks") or []) + lr.hyperlinks
                if lr.note_ids:
                    delim["note_ids"] = sorted(set((delim.get("note_ids") or []) + lr.note_ids))
                state = "preamble"
                i += 1
                continue
            state = "preamble"  # no distinct title; reparse this row below

        # 8. titles (front region) --------------------------------------------
        if state == "front" or not seen_opening:
            title = _match_title(c, lr)
            if title is not None:
                ttype, prefix, ttext = title
                # sem-v5: a short decision prints title and decision on one line;
                # emit the decision as its own body paragraph instead of burying it
                # in the title element.
                ttext, body_text = _split_title_body(ttext)
                elements.append(_new_element(lr, type="title", section=section,
                                             prefix=prefix, text=ttext,
                                             text_index=text_index))
                seen_title = True
                if body_text:
                    body_el = _new_element(lr, type="paragraph", section=section,
                                           paragraph_type=None, level=1,
                                           text=body_text, text_index=text_index)
                    body_el["split_continuation"] = True
                    elements.append(body_el)
                    state = "decision"      # the body has started; never frontmatter
                # PRSTs have no opening formula: their quoted body is a statement,
                # so switch out of the frontmatter phase here.
                if TITLE_PRST_RE.match(c):
                    state = "statement"
                i += 1
                continue

        # 9. vote record: consume the whole block into one element -----------
        if VOTE_RE.match(c) or VOTE_TALLY_RE.match(c):
            j, positions, summary, vote, vsum = _consume_vote_block(lrows, i)
            el = {
                "positions": positions,
                "type": "vote_record",
                "section": section,
                "paragraph_type": None,
                "level": None,
                "prefix": None,
                "heading_level": None,
                "text": summary,
                "lead_verb": None,
                "hyperlinks": [],
                "note_ids": [],
                "vote": vote,
            }
            if vsum:
                el["vote_summary"] = vsum
            if text_index != 1:
                el["text_index"] = text_index
            elements.append(el)
            state = "tail"
            i = j if j > i else i + 1
            continue

        # 10. signature / meeting line ----------------------------------------
        if MEETING_RE.match(c) or (state == "tail" and DATE_LINE_RE.match(c)):
            elements.append(_new_element(lr, type="signature", section=section,
                                         text=c, text_index=text_index))
            state = "tail"
            i += 1
            continue

        # 11. operative paragraph ---------------------------------------------
        # NB: paragraph_type='operative'/'preambular' is a *main-section* concept.
        # Annex/appendix bodies are frequently numbered too (programmes of action,
        # agendas), but those are backmatter, not resolution operatives -- so we
        # keep their prefix/level yet set paragraph_type=None there.
        # label operatives in the main section AND inside a scoped instrument annex
        in_main = section == "main"
        label_ops = in_main or annex_scoped
        # sem-v5: an unconfirmed leading number is NOT a marker. The row falls
        # through to the ordinary text paths below with its number still inside the
        # text, so nothing is deleted and nothing is invented.
        m_num = OP_NUM_RE.match(c) if i in op_numbers else None
        if m_num and not TITLE_GA_NUM_RE.match(c) and not _style_is_heading(c, lr, state):
            op_tracker.top()
            last_op_number = int(m_num.group(1))
            elements.append(_new_element(
                lr, type="paragraph", section=section,
                paragraph_type="operative" if label_ops else None,
                level=1, prefix=f"{m_num.group(1)}.", text=m_num.group(2).strip(),
                lead_verb=_op_lead_verb(m_num.group(2)), text_index=text_index))
            state = "operative"
            i += 1
            continue

        # tolerant operative number (missing/misplaced period), gated hard:
        # only continue an already-running operative sequence, and only when the
        # clause opens with a finite operative verb -- so "232 .Recognizes",
        # "4 Calls", "13 Urges" are rescued without misreading years/quantities.
        # sem-v5: the number must also CONTINUE the sequence (last+1..last+gap, or
        # a restart at 1); a number the confirmation pass already refused can never
        # come back in through this path.
        if (state == "operative" and not m_num and not TITLE_GA_NUM_RE.match(c)
                and i not in refused_idx):
            m_loose = OP_NUM_LOOSE_RE.match(c)
            if m_loose and (int(m_loose.group(1)) == 1
                            or 0 < int(m_loose.group(1)) - last_op_number <= MARKER_MAX_GAP):
                rest = m_loose.group(2).strip()
                w0 = rest.split(" ", 1)[0].strip(",.").lower()
                # source may drop the period AND glue the number to an adverb-led
                # verb ("6Also reaffirms ..."): look past a leading adverb for the
                # finite operative verb before deciding.
                verb = w0
                if w0 in ("also", "further", "again", "finally", "moreover") and " " in rest:
                    verb = rest.split(" ", 2)[1].strip(",.").lower()
                if verb in OPERATIVE_LEAD_VERBS:
                    op_tracker.top()
                    last_op_number = int(m_loose.group(1))
                    elements.append(_new_element(
                        lr, type="paragraph", section=section,
                        paragraph_type="operative" if label_ops else None,
                        level=1, prefix=f"{m_loose.group(1)}.",
                        text=rest,
                        lead_verb=_op_lead_verb(rest), text_index=text_index))
                    state = "operative"
                    i += 1
                    continue

        # sem-v5: an upper-case Roman parenthetical is only a subparagraph marker
        # when the sequence confirms it; otherwise it is a citation tail
        # ('… resolution 900 (IX) of 11 December 1954') and stays inside the text.
        m_par = OP_PAREN_RE.match(c)
        if (m_par and state in ("operative", "preamble")
                and _UPPER_ROMAN_RE.match(m_par.group(1))
                and i not in upper_roman_parens):
            refused_markers.append((i, 0, c[:60]))
            m_par = None
        if m_par and state in ("operative", "preamble"):
            level, prefix = op_tracker.classify_paren(
                m_par.group(1), next_tok=_next_paren_token(lrows, i + 1))
            # A parenthetical in the PREAMBLE is a sub-item of the preceding
            # preambular clause (often introduced by a clause ending in ':'),
            # NOT the first operative -- operatives are introduced by "1." So we
            # keep it preambular and do NOT switch state. Only in the operative
            # part is a parenthetical an operative subparagraph.
            base = "operative" if state == "operative" else "preambular"
            elements.append(_new_element(
                lr, type="paragraph", section=section,
                paragraph_type=base if label_ops else None,
                level=level, prefix=prefix, text=m_par.group(2).strip(),
                text_index=text_index))
            i += 1
            continue

        # 11c. dotted enumerator (sem-v4): a confirmed 'N.N'/'N.a' sub-section marker
        # carried INSIDE the paragraph text (SDG targets under 'Goal N.', forest-plan
        # targets, programme-plan '19.1' paragraphs). The pre-pass has verified this
        # row sits in a coherent contiguous run, so the marker moves to `prefix`, the
        # clause gets level=2, and the text is stripped. paragraph_type stays as it
        # would otherwise be (None here — these are sub-heading content, not numbered
        # resolution operatives), so operative counts are not inflated.
        if i in dotted_enum:
            d_prefix, d_level, d_rest = dotted_enum[i]
            elements.append(_new_element(
                lr, type="paragraph", section=section, paragraph_type=None,
                level=d_level, prefix=d_prefix, text=d_rest, text_index=text_index))
            i += 1
            continue

        # 12. body heading (roman/letter/short heading-styled/bold-run-in line) --
        # sem-v3: split any leading section marker into `prefix`; a style heading
        # takes its tier from the style, a bold run-in label nests one level below
        # the enclosing section heading.
        if _looks_like_heading(c, lr, state):
            hpref, htext = _split_heading_prefix(c)
            if _style_is_heading(c, lr, state):
                hlevel = _heading_level(lr)
                last_heading_level = hlevel
            elif _bold_heading(c, lr, state):
                hlevel = min((last_heading_level or 2) + 1, 6)
            else:
                hlevel = _heading_level(lr)
                last_heading_level = hlevel
            elements.append(_new_element(lr, type="heading", section=section,
                                         heading_level=hlevel, prefix=hpref,
                                         text=htext, text_index=text_index))
            i += 1
            continue

        # 12b. sem-v5 EXIT FROM THE FRONT REGION.
        # Rule 13 below is a catch-all: while the machine is in `front` EVERY line
        # becomes type='frontmatter', which the site hides. Decisions have no
        # opening formula, and a PDF/OCR resolution whose formula is garbled
        # ('Thc General .111·n11h/t·.') never matches OPENING_RE either -- so whole
        # bodies were being hidden (the blank-page class). Two shapes end the front
        # region here, both body text by any reading:
        #   * the decision formula ('At its 87th plenary meeting, on 30 June 2023,
        #     the General Assembly … decided …');
        #   * a substantial clause opening with a finite operative verb or a
        #     preambular participle ('Recalling its resolution 1514 (XV) …').
        # Both emit a real paragraph; paragraph_type stays NULL so no clause is
        # labelled operative/preambular on this evidence alone.
        if state in ("front", "blockhead") and not _heading_fp(c):
            first = c.split(" ", 1)[0].strip(",.").lower()
            body_shaped = (len(c.split()) >= 8
                           and (first in OPERATIVE_LEAD_VERBS
                                or first in PREAMBULAR_FIRST_WORDS))
            # a row whose leading number the sequence pass REFUSED is still body
            # text; refusing a marker must never move content into the hidden
            # frontmatter bucket (it would trade a fabrication for a blank page).
            if (DECISION_OPENING_RE.match(c) or body_shaped
                    or i in refused_idx or OP_NUM_RE.match(c)):
                elements.append(_new_element(
                    lr, type="paragraph", section=section, paragraph_type=None,
                    level=1, text=c, text_index=text_index))
                state = "decision"
                i += 1
                continue

        # 13. frontmatter residue (session/agenda/masthead lines) -------------
        if state in ("front", "blockhead"):
            # drop obvious page artifacts, keep informative masthead as frontmatter
            if RUNNING_HEADER_RE.match(c) or PAGE_NUM_RE.match(c):
                for p in lr.positions:
                    dropped.append({"position": p, "reason": "page_artifact"})
            else:
                st = "masthead" if (MASTHEAD_RE.match(c) or DATE_LINE_RE.match(c)) else None
                elements.append(_new_element(lr, type="frontmatter", section=section,
                                             text=c, text_index=text_index, subtype=st))
            i += 1
            continue

        # 14. preambular ------------------------------------------------------
        if state == "preamble":
            lead_it = lr.props.get("lead_italic_text")
            first = c.split(" ", 1)[0].strip(",").lower()
            # Unnumbered operative: some short resolutions have a single operative
            # clause with no "1." -- it is a *finite* verb (Decides/Requests...),
            # whereas preambular leads are participles (-ing) or adjectives.
            lead_word = (lead_it or c).split(" ", 1)[0].strip(",").lower()
            if lead_word in OPERATIVE_LEAD_VERBS and first not in PREAMBULAR_FIRST_WORDS:
                op_tracker.top()
                elements.append(_new_element(
                    lr, type="paragraph", section=section,
                    paragraph_type="operative" if label_ops else None,
                    level=1, prefix=None, text=c, lead_verb=_op_lead_verb(c),
                    text_index=text_index))
                state = "operative"
                i += 1
                continue
            lead = lead_it or _lead_verb_from_text(c)
            elements.append(_new_element(
                lr, type="paragraph", section=section,
                paragraph_type="preambular" if label_ops else None,
                level=1, text=c, lead_verb=lead, text_index=text_index))
            i += 1
            continue

        # 14b. OPT-IN source-defect rescue: an unlabeled operative whose NUMBER was
        # dropped at source. Fires only inside a running operative sequence with a
        # CONFIRMED numbering gap ahead, and only for clauses that read operative.
        if (RESCUE_INFERRED_OPERATIVE and state == "operative" and label_ops
                and not OP_NUM_RE.match(c) and not OP_PAREN_RE.match(c)
                and not _looks_like_heading(c, lr, state)):
            first = c.split(" ", 1)[0].strip(",.").lower()
            in_alpha_run = op_tracker.last_alpha is not None and not op_tracker.roman_active
            looks_op = first in OPERATIVE_LEAD_VERBS or INFINITIVE_SUBITEM_RE.match(c)
            inferred = False
            level = 1
            if looks_op:
                if in_alpha_run:
                    nxt = _next_number_ahead(lrows, i + 1, paren=True)
                    cur_ord = _alpha_ord(op_tracker.last_alpha) or 0
                    if nxt is not None and nxt > cur_ord + 1:
                        inferred, level = True, 2
                else:
                    # sem-v5: the "confirmed gap ahead" that licenses the rescue must
                    # be a CONFIRMED marker -- an invented number ahead must never
                    # license inventing an operative label here.
                    nxt = _next_number_ahead(lrows, i + 1, paren=False,
                                             confirmed=op_numbers)
                    if nxt is not None and nxt > last_op_number + 1:
                        inferred, level = True, 1
            if inferred:
                el = _new_element(
                    lr, type="paragraph", section=section,
                    paragraph_type="operative", level=level, prefix=None, text=c,
                    lead_verb=_op_lead_verb(c) if level == 1 else None,
                    text_index=text_index)
                el["inferred_operative"] = True
                elements.append(el)
                i += 1
                continue

        # 15. operative continuation / tail body / annex prose / PRST body ----
        if state in ("operative", "tail"):
            # An UNPREFIXED line inside the operative part is a continuation /
            # chapeau / stray heading, not a distinct numbered operative -- a
            # real operative clause always carries a prefix (or was caught as an
            # unnumbered operative in step 14). So paragraph_type=None here, to
            # avoid inflating operative counts.
            elements.append(_new_element(
                lr, type="paragraph", section=section, paragraph_type=None,
                level=None, text=c, text_index=text_index))
            i += 1
            continue

        # 16. PRST / statement body (no opening formula seen) -----------------
        if not seen_opening:
            elements.append(_new_element(lr, type="paragraph", section=section,
                                         paragraph_type=None, level=1, text=c,
                                         text_index=text_index))
            i += 1
            continue

        # 17. fallback: unclassified -----------------------------------------
        issues.append({"position": lr.positions[0], "problem": "unclassified paragraph",
                       "text_head": c[:80]})
        elements.append(_new_element(lr, type="paragraph", section=section,
                                     paragraph_type=None, text=c, text_index=text_index))
        i += 1

    # sem-v3: fold a bare section-marker heading ('I') into the short title line
    # that follows it, then backfill annex_index onto every annexed element. Both
    # are position-preserving (concatenate/copy only), so accounting is unaffected.
    elements = _merge_bare_headings(elements)
    _stamp_annex_index(elements)

    # sem-v2: annotate operative/preambular elements with their action verb in a
    # cleanly separable pass (additive; never touches text/positions/accounting).
    annotate_actions(elements)

    # sem-v5: every refused marker is LEDGERED (a drop with no ledger is
    # unauditable). The raw position is recorded so the claim can be checked
    # against the source; `issues[]` is persisted verbatim to
    # digitallibrary.document_parses, so the refusals are queryable in SQL.
    for lidx, num, head in sorted(refused_markers):
        pos = lrows[lidx].positions[0] if lrows[lidx].positions else -1
        issues.append({
            "position": pos,
            "problem": "unconfirmed_marker",
            "text_head": (f"{num}. " if num else "") + head,
        })

    result = {
        "symbol": symbol,
        "format": fmt,
        "parser_version": PARSER_VERSION,
        "elements": elements,
        "dropped": dropped,
        "issues": issues,
    }
    return result


def _alpha_ord(s: str) -> int | None:
    """Ordinal of a subparagraph letter, honoring the UN doubling convention past z
    ('aa'=27,'bb'=28,..). Mirrors fulltext_review._alpha_value so gap arithmetic
    agrees on both sides."""
    s = s.lower()
    if not s.isalpha():
        return None
    if len(s) >= 2 and len(set(s)) == 1:
        return (len(s) - 1) * 26 + (ord(s[0]) - ord("a") + 1)
    if len(s) == 1:
        return ord(s) - ord("a") + 1
    return None


def _next_number_ahead(lrows: list[LRow], i: int, paren: bool,
                       confirmed: dict[int, int] | None = None) -> int | None:
    """Ordinal of the next labeled operative item ahead of lrows[i], or None.

    paren=False -> next top-level 'N.' number; paren=True -> next '(letter)' alpha
    ordinal. Scans a bounded window and stops at a hard structural boundary so the
    look-ahead never crosses into another text/annex/preamble. When `confirmed` is
    given (sem-v5), only numbers the sequence pass accepted count as a number ahead.
    """
    n = len(lrows)
    j = i
    steps = 0
    while j < n and steps < 30:
        cj = lrows[j].clean
        if not cj or lrows[j].kind in ("empty", "section_break", "footnote"):
            j += 1
            continue
        if (OPENING_RE.match(cj) or ANNEX_RE.match(cj) or MEETING_RE.match(cj)
                or VOTE_RE.match(cj) or VOTE_TALLY_RE.match(cj) or DIVIDER_RE.match(cj)):
            return None
        if not paren:
            m = OP_NUM_RE.match(cj)
            if m and not TITLE_GA_NUM_RE.match(cj):
                if confirmed is not None and j not in confirmed:
                    j += 1
                    steps += 1
                    continue
                return int(m.group(1))
        else:
            m = OP_PAREN_RE.match(cj)
            if m:
                return _alpha_ord(m.group(1))
        j += 1
        steps += 1
    return None


def _next_paren_token(lrows: list[LRow], i: int) -> str | None:
    """The next '(marker)' token at/after lrows[i] within a bounded window, or None.
    Stops at a hard structural boundary so it never crosses into another run."""
    n = len(lrows)
    j, steps = i, 0
    while j < n and steps < 20:
        cj = lrows[j].clean
        if not cj or lrows[j].kind in ("empty", "section_break", "footnote"):
            j += 1
            continue
        if (OPENING_RE.match(cj) or ANNEX_RE.match(cj) or MEETING_RE.match(cj)
                or VOTE_RE.match(cj) or DIVIDER_RE.match(cj) or OP_NUM_RE.match(cj)):
            return None
        m = OP_PAREN_RE.match(cj)
        if m:
            return m.group(1)
        j += 1
        steps += 1
    return None


def _op_lead_verb(rest: str) -> str | None:
    """Leading verb of an operative clause (first word, if title-case-ish)."""
    words = rest.strip().split()
    if not words:
        return None
    w = words[0].strip(",")
    if w[:1].isupper() and w.isalpha():
        # "Also welcomes" / "Further requests" -> include the adverb
        if w.lower() in ("also", "further", "again", "finally", "moreover") and len(words) > 1:
            return f"{w} {words[1].strip(',')}"
        return w
    return None


def _split_title_body(text: str) -> tuple[str, str | None]:
    """Cut a run-together decision line into (title, body).

    Short decisions print the title and the decision itself on ONE line:
    '77/408. Appointment of members of the ACABQ At its 87th plenary meeting, on
    30 June 2023, the General Assembly … appointed …'. Everything used to become
    one type='title' element (479 documents, median 326 characters), so the
    decision itself was never rendered as body text. The cut is made at the
    sentence boundary immediately before the decision formula; if the formula opens
    the line there is nothing to split.
    """
    m = DECISION_OPENING_RE.search(text)
    if not m or m.start() == 0:
        return text, None
    head = text[:m.start()].rstrip().rstrip(",;:").rstrip()
    body = text[m.start():].strip()
    if not head or len(body.split()) < 4:
        return text, None
    return head, body


def _match_title(c: str, lr: LRow) -> tuple[str, str | None, str] | None:
    """Return (type, prefix, text) if `c` is a resolution/PRST title, else None."""
    if TITLE_PRST_RE.match(c):
        return ("title", None, c)
    if TITLE_SC_RE.match(c):
        return ("title", None, c)
    m = TITLE_GA_NUM_RE.match(c)
    if m:
        return ("title", f"{m.group(1)}.", m.group(2).strip())
    # HRC "15/7 Subject" -- require a heading style to avoid matching cross refs
    if lr.style in ("HChG", "HCh", "TitleH1", "TitleHCH") and TITLE_HRC_NUM_RE.match(c):
        m2 = TITLE_HRC_NUM_RE.match(c)
        return ("title", m2.group(1), m2.group(2).strip())
    return None


# A lettered/roman section heading with a descriptive title ('A. Mandate',
# 'B. Advancing integration ...', 'III. Green economy ...'). Split so we can read
# the title's first word.
_MARKER_HEAD_RE = re.compile(r"^((?:[IVXLC]{1,4}|[A-Z])\.)\s+([A-Z]\S.*)$")


def _marker_section_heading(c: str, lr: LRow) -> bool:
    """sem-v4: a bold / heading-styled lettered-or-roman section heading whose title
    may be LONG ('B. Advancing integration, implementation and coherence: ...',
    'III. Green economy in the context of ...'). Requires bold OR an explicit heading
    style; the title must NOT open with a finite operative/participial lead verb, so a
    bold 'I. Decides that ...' / 'II. Recalling ...' stays a clause, never a heading."""
    if not (lr.props.get("bold") or lr.style in BODY_HEADING_STYLES):
        return False
    m = _MARKER_HEAD_RE.match(c)
    if not m:
        return False
    w0 = m.group(2).split(" ", 1)[0].strip(",.").lower()
    return w0 not in OPERATIVE_LEAD_VERBS and w0 not in PREAMBULAR_FIRST_WORDS


def _looks_like_heading(c: str, lr: LRow, state: str) -> bool:
    """Heuristic for a section heading inside the body (not title/opening)."""
    if state == "front":
        return False
    # sem-v3: trust the document's own heading style / bold run-in labels first
    # (these may be long, so they precede the length guard).
    if _style_is_heading(c, lr, state) or _bold_heading(c, lr, state):
        return True
    # sem-v4: a bold/styled lettered-or-roman section heading with a (possibly long)
    # descriptive title, guarded so bold operative clauses ('I. Decides ...') are not
    # swallowed -- also precedes the length guard.
    if _marker_section_heading(c, lr):
        return True
    if len(c) > 80:
        return False
    # lone roman numeral or single capital letter (consolidated sub-res letter,
    # or numbered section heading within the operative part)
    if re.match(r"^[IVXLC]{1,4}\.?$", c) or re.match(r"^[A-Z]\.?$", c):
        return True
    # explicit heading style, short, and centered/bold
    if lr.style in BODY_HEADING_STYLES and (lr.props.get("bold") or lr.props.get("alignment") == "center"):
        return True
    return False


def _heading_level(lr: LRow) -> int:
    """Heading tier from the paragraph style, mirroring the TOC verifier: H<n>/H<nm>
    -> first digit (H23 -> 2), Heading<n> -> that digit, HCh/unknown -> 1."""
    st = lr.style or ""
    m = re.match(r"^H(\d)", st)
    if m:
        return int(m.group(1))
    if STYLE_HEADING_RE.match(st):
        d = re.search(r"\d", st)
        if d:
            return int(d.group())
    return 1


# ---------------------------------------------------------------------------
# I/O
# ---------------------------------------------------------------------------


def fetch_targets(limit: int | None, symbol: str | None,
                  offset: int = 0) -> list[tuple[str, str, str, str | None, str | None]]:
    """Return (symbol_normalized, lang, format, archive_path, converted_path).

    Targets any doc whose raw extraction is available: status IN
    ('extracted', 'parsed'). Including 'parsed' keeps re-parses working after the
    loader has advanced status (a plain JSON re-run, or a re-load with --to-db,
    still finds every already-loaded doc).
    """
    sql = (
        "SELECT df.symbol_normalized, df.lang, df.format, df.archive_path, df.converted_path "
        "FROM digitallibrary.document_files df "
        "WHERE df.status IN ('extracted', 'parsed') "
    )
    params: list[object] = []
    if symbol:
        sql += "AND df.symbol_normalized = %s "
        params.append(symbol)
    sql += "ORDER BY (df.status = 'extracted') DESC, df.symbol_normalized "
    if limit:
        sql += "LIMIT %s "
        params.append(limit)
    if offset:
        sql += "OFFSET %s"
        params.append(offset)
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(sql, params)
        return [(r[0], r[1] or "en", r[2], r[3], r[4]) for r in cur.fetchall()]


def fetch_rows(conn, symbol: str, lang: str = "en") -> list[dict]:
    with conn.cursor() as cur:
        cur.execute(
            "SELECT position, kind, text, style_id, style_name, numbering, props, "
            "table_cell, hyperlinks, footnote_ref "
            "FROM digitallibrary.document_paragraphs_raw "
            "WHERE symbol_normalized = %s AND lang = %s ORDER BY position",
            [symbol, lang],
        )
        cols = [d[0] for d in cur.description]
        return [dict(zip(cols, row)) for row in cur.fetchall()]


# ---------------------------------------------------------------------------
# Semantic DB loader (migration 003 tables)
# ---------------------------------------------------------------------------

# Insert column order for digitallibrary.document_paragraphs. `id` and `position`
# are computed by the loader; the rest are read off each parsed element.
_PARA_COLUMNS = (
    "symbol_normalized", "lang", "position", "id", "type", "subtype", "section",
    "annex_index", "text_index", "paragraph_type", "level", "heading_level",
    "prefix", "lead_verb", "text", "raw_positions", "inferred_operative",
    "vote", "vote_summary", "hyperlinks", "note_ids", "parser_version",
    # migration 004: flattened action-verb annotation (NULL unless the element
    # carries a nested `action` object, i.e. an annotated operative/preambular clause)
    "action_verb", "action_verb_normalized", "action_category", "action_force",
    "action_sentiment", "action_bindingness", "action_budget_relevant",
    "action_modifiers", "assignee", "assignee_head_noun", "assignee_class",
    "action_inherited", "action_context_marker",
)
_PARA_INSERT = (
    f"INSERT INTO digitallibrary.document_paragraphs ({', '.join(_PARA_COLUMNS)}) "
    f"VALUES ({', '.join(['%s'] * len(_PARA_COLUMNS))})"
)


def element_uuid(symbol: str, lang: str, position: int) -> uuid.UUID:
    """Deterministic element id: uuid5(NAMESPACE_URL, '<symbol>:<lang>:<position>').

    `position` is the 0-based element index in parsed order. Stable across
    re-parses as long as element ordering is stable.
    """
    return uuid.uuid5(uuid.NAMESPACE_URL, f"{symbol}:{lang}:{position}")


def load_document(conn, symbol: str, lang: str, fmt: str, result: dict) -> int:
    """Delete-then-insert one parsed document into the semantic tables.

    Writes document_paragraphs (one row per element, 0-based position) and one
    document_parses ledger row (element_count + JSON root dropped[]/issues[]).
    Caller owns the transaction (commit/rollback per doc). Returns element count.
    """
    elements = result["elements"]
    with conn.cursor() as cur:
        cur.execute(
            "DELETE FROM digitallibrary.document_paragraphs "
            "WHERE symbol_normalized = %s AND lang = %s",
            [symbol, lang],
        )
        cur.execute(
            "DELETE FROM digitallibrary.document_parses "
            "WHERE symbol_normalized = %s AND lang = %s",
            [symbol, lang],
        )
        rows = []
        for pos, el in enumerate(elements):
            vote = el.get("vote")
            vote_summary = el.get("vote_summary")
            action = el.get("action")
            assignee = action.get("assignee") if action else None
            rows.append((
                symbol, lang, pos, str(element_uuid(symbol, lang, pos)),
                el["type"], el.get("subtype"), el.get("section", "main"),
                el.get("annex_index"), el.get("text_index", 1),
                el.get("paragraph_type"), el.get("level"), el.get("heading_level"),
                el.get("prefix"), el.get("lead_verb"), el["text"],
                el["positions"], bool(el.get("inferred_operative", False)),
                json.dumps(vote) if vote is not None else None,
                json.dumps(vote_summary) if vote_summary is not None else None,
                json.dumps(el.get("hyperlinks") or []),
                json.dumps(el.get("note_ids") or []),
                result["parser_version"],
                # migration 004: flattened action annotation (all NULL when no action)
                action.get("verb") if action else None,
                action.get("normalized") if action else None,
                action.get("category") if action else None,
                action.get("force") if action else None,
                action.get("sentiment") if action else None,
                action.get("bindingness") if action else None,
                action.get("budget_relevant") if action else None,
                json.dumps(action["modifiers"]) if action and action.get("modifiers") else None,
                assignee.get("verbatim") if assignee else None,
                assignee.get("head_noun") if assignee else None,
                assignee.get("addressee_class") if assignee else None,
                action.get("inherited") if action else None,
                action.get("context_marker") if action else None,
            ))
        if rows:
            cur.executemany(_PARA_INSERT, rows)
        cur.execute(
            "INSERT INTO digitallibrary.document_parses "
            "(symbol_normalized, lang, parser_version, format, element_count, dropped, issues) "
            "VALUES (%s, %s, %s, %s, %s, %s::jsonb, %s::jsonb)",
            [symbol, lang, result["parser_version"], fmt, len(elements),
             json.dumps(result.get("dropped", [])), json.dumps(result.get("issues", []))],
        )
    return len(elements)


# ---------------------------------------------------------------------------
# sem-v5: CONSERVATION CHECKS
#
# Two denominators, neither of them derived from what the parse produced:
#   (1) word accounting  -- every WORD of every raw row must end up in an element
#       (text, prefix, structured vote lists, footnote marker ids) or in a raw row
#       that `dropped[]` names with a reason. The pre-existing invariant only
#       asserted that every raw POSITION was consumed by SOME element, which stays
#       true when an element quietly loses text.
#   (2) source conservation -- the archived SOURCE FILE is re-read here, by this
#       module, and the words the parse holds are compared against it. For a Word
#       source that is the whole file; for a PDF page holding several documents it
#       is the target document's own span (its number/title anchor up to its
#       adoption record or the next document's heading). Below TRUNCATION_FLOOR the
#       document FAILS: it is not stored as 'parsed', the run exits non-zero, and
#       the reason is written to the ledger. No extractor flag can suppress it --
#       the 145 silently truncated documents all carried `anchor_found=True` and an
#       empty flag list.
# ---------------------------------------------------------------------------

_WORD_RE = re.compile(r"[0-9A-Za-z]+")

# Bar for source conservation. Word-format sources sit at >=0.85 (p01 of a
# 150-document sample; minimum observed 0.76) and PDF span ratios at >=0.67 (p05 of
# a 181-document sample, median 1.03), while the known-truncated documents sit at
# 0.015-0.28 -- so 0.50 separates them with headroom on both sides and is NOT the
# constraint that healthy documents scrape past.
TRUNCATION_FLOOR = 0.50
# Fallback floor, used when the span anchor cannot be trusted: the parse against the
# AVERAGE document on its own source page (page words / documents the page prints,
# counted from the page's own headings). A document holding less than a quarter of
# that is truncated. This is what catches a truncation whose span anchor landed in
# the page's table of contents (`A/RES/32/1`: 10 words where the page average is 56).
WHOLE_SOURCE_FLOOR = 0.25
# A span shorter than this is not a document, it is a table-of-contents entry or a
# running head: the anchor is untrustworthy and the whole-source floor decides
# instead. (The shortest genuine span seen in a 1,200-document sample is 42 words.)
MIN_TRUSTED_SPAN_WORDS = 40

# The organ opening formula as it appears in a SOURCE file (looser than OPENING_RE,
# which matches a whole raw row: in a source read the formula sits inside a line).
_SOURCE_OPENING_RE = re.compile(
    r"\bThe\s+(General\s*,?\s*Assembly|Security\s+Council|Economic\s+and\s+Social\s+Council|"
    r"Human\s+Rights\s+Council|Trusteeship\s+Council)\s*[,.]", re.I)

# Old-style volume heading ("1110 (XI). Admission of …") and modern numbered title
# ("32/1. Question of …"): where one document ends and the next begins on a shared
# source page.
_SRC_HEAD_OLD_RE = re.compile(r"(?m)^\s*\d{1,4}\s*\([IVXLCDM\-]+\)\s*\.")
_SRC_HEAD_NEW_RE = re.compile(r"(?m)^\s*\d{1,3}/\d{1,4}[A-Z]?\s*\.")
_SRC_HEAD_RESOLUTION_RE = re.compile(r"(?mi)^\s*Resolution\s+\d{1,4}\s*\(")
# The adoption record that closes a resolution ("425th plenary meeting", "Adopted at
# the 2932nd meeting", "Adopted without a vote").
_SRC_ADOPTION_RE = re.compile(
    r"\b\d{1,4}\s*(?:st|nd|rd|th|d)\b[^\n]{0,24}\bmeeting\b"
    r"|\bAdopted\s+(?:at|by|unanimously|without)\b", re.I)

# Symbol shapes we can anchor inside a source page.
_SYM_OLD_RE = re.compile(r"^[AES]/(?:RES|DEC)/(\d{1,4})\(([A-Z][A-Z0-9\-]*)\)$")
_SYM_NEW_RE = re.compile(r"^[AE]/(?:RES|DEC)/(\d{1,4})/(\d{1,4})[A-Z\-]*$")
_SYM_SC_RE = re.compile(r"^S/RES/(\d{1,4})\((\d{4})\)$")


# A heading belonging to ANOTHER document. Truncation is not the only crop defect:
# a crop that runs PAST the document's own printed extent stores the neighbouring
# resolution's text under this symbol, which is fabrication rather than loss
# (`A/RES/1005(ES-II)` carries its neighbour's entire preamble). These match at the
# START of an element only, so an in-sentence citation ('… resolution 1514 (XV) of
# 14 December 1960 …') is never mistaken for a heading.
_FOREIGN_HEAD_OLD_RE = re.compile(r"^(\d{1,4})\s*\(([A-Z][A-Z0-9\-]*)\)\s*[.,]")
_FOREIGN_HEAD_NEW_RE = re.compile(r"^(\d{1,3}/\d{1,4})[A-Z]?\s*\.")
_FOREIGN_HEAD_SC_RE = re.compile(r"^Resolution\s+(\d{1,4})\s*\(", re.I)


def _own_document_numbers(symbol: str) -> set[str]:
    """The document's own printed number(s), as they appear in a heading."""
    out: set[str] = set()
    m = _SYM_OLD_RE.match(symbol) or _SYM_SC_RE.match(symbol)
    if m:
        out.add(m.group(1))
    m = _SYM_NEW_RE.match(symbol)
    if m:
        out.add(f"{m.group(1)}/{m.group(2)}")
        out.add(m.group(2))
    return out


def check_overreach(symbol: str, result: dict) -> list[str]:
    """Headings of OTHER documents inside this parse (the crop ran past its extent).

    Returns the foreign document numbers found, in document order."""
    own = _own_document_numbers(symbol)
    if not own:
        return []
    found: list[str] = []
    for el in result["elements"]:
        # Footnotes ARE citations of other resolutions ('Resolution 1514 (XV).');
        # only body/heading/title text can carry a neighbour's heading.
        if el.get("type") in ("footnote", "vote_record", "divider", "table",
                              "signature", "backmatter"):
            continue
        text = ((el.get("prefix") or "") + " " + (el.get("text") or "")).strip()
        for rx in (_FOREIGN_HEAD_OLD_RE, _FOREIGN_HEAD_NEW_RE, _FOREIGN_HEAD_SC_RE):
            m = rx.match(text)
            if not m:
                continue
            num = m.group(1)
            if num not in own and num.split("/")[-1] not in own and num not in found:
                found.append(num)
            break
    return found


def _word_count(text: str) -> int:
    return len(_WORD_RE.findall(text or ""))


def _word_bag(text: str) -> dict[str, int]:
    bag: dict[str, int] = {}
    for w in _WORD_RE.findall(text or ""):
        w = w.lower()
        bag[w] = bag.get(w, 0) + 1
    return bag


def parse_word_count(result: dict) -> int:
    """Words the parse actually holds (element prefixes + text + vote country lists)."""
    n = 0
    for el in result["elements"]:
        n += _word_count((el.get("prefix") or "") + " " + (el.get("text") or ""))
        vote = el.get("vote")
        if vote:
            for lst in vote.values():
                n += _word_count(" ".join(lst))
    return n


def check_word_accounting(result: dict, raw_rows: list[dict]) -> list[str]:
    """Words of the raw rows that no element and no ledgered drop accounts for.

    Returned as a sorted 'token xN' list (empty = conserved). The only structural
    re-encodings allowed are named here: vote tally LABELS become the `vote` dict's
    keys, and a footnote's leading marker glyph becomes `note_ids`.
    """
    raw_bag: dict[str, int] = {}
    for row in raw_rows:
        for w, c in _word_bag(row.get("text") or "").items():
            raw_bag[w] = raw_bag.get(w, 0) + c
    out_bag: dict[str, int] = {}

    def add(text: str) -> None:
        for w, c in _word_bag(text).items():
            out_bag[w] = out_bag.get(w, 0) + c

    for el in result["elements"]:
        add((el.get("prefix") or "") + " " + (el.get("text") or ""))
        vote = el.get("vote")
        if vote:
            for key, lst in vote.items():
                add(" ".join(lst))
                add(key.replace("_", " "))          # 'In favour:' label -> vote key
        if el.get("type") == "footnote":
            add(" ".join(str(x) for x in (el.get("note_ids") or [])))   # marker glyph
    by_pos = {row["position"]: row.get("text") or "" for row in raw_rows}
    for d in result.get("dropped", []):
        add(by_pos.get(d["position"], ""))

    missing = []
    for w, c in raw_bag.items():
        gap = c - out_bag.get(w, 0)
        if gap > 0:
            missing.append((w, gap))
    missing.sort(key=lambda x: (-x[1], x[0]))
    return [f"{w} x{c}" if c > 1 else w for w, c in missing]


def read_source_text(fmt: str, archive_path: str | None,
                     converted_path: str | None) -> str:
    """Re-read the ARCHIVED SOURCE FILE independently of the extraction layer.

    Raises FileNotFoundError / any reader error to the caller: an unreadable source
    is reported as a failure, never as a silent pass.
    """
    if fmt == "pdf":
        rel = archive_path
    else:
        rel = converted_path or archive_path
    if not rel:
        raise FileNotFoundError("no archive_path on the ledger row")
    path = ARCHIVE_ROOT / rel
    if not path.exists():
        raise FileNotFoundError(str(path))
    if path.suffix.lower() == ".pdf":
        import fitz                                     # pymupdf
        with fitz.open(str(path)) as doc:
            return "\n".join(page.get_text() for page in doc)
    if path.suffix.lower() == ".docx":
        import docx                                     # python-docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    return path.read_text(errors="ignore")


def _source_anchor(symbol: str) -> re.Pattern | None:
    """Regex locating the target document's own number inside a source page."""
    m = _SYM_OLD_RE.match(symbol)
    if m:
        return re.compile(rf"{m.group(1)}\s*\(\s*{m.group(2)}\s*\)")
    m = _SYM_SC_RE.match(symbol)
    if m:
        return re.compile(rf"{m.group(1)}\s*\(\s*{m.group(2)}\s*\)")
    m = _SYM_NEW_RE.match(symbol)
    if m:
        return re.compile(rf"\b{m.group(1)}\s*/\s*{m.group(2)}\b")
    return None


def source_span_words(symbol: str, src: str) -> int | None:
    """Words of the target document's own span inside a multi-document source page.

    From the document's number anchor to whichever comes first: its adoption record
    or the next document's heading. The anchor usually occurs several times on a
    page (the printed heading, running heads, cross-references such as 'See
    resolution 357 (IV) on this page'), so occurrences are ranked: a HEADING-shaped
    one (line start, followed by a period) that contains the organ opening formula
    wins, then heading-shaped, then any occurrence whose span holds the opening
    formula, then the rest. Within a rank the SMALLEST span is taken, so a duplicate
    heading never inflates the denominator. None when the symbol shape cannot be
    anchored -- reported as 'unchecked', never as a pass.
    """
    rx = _source_anchor(symbol)
    if rx is None:
        return None
    ranked: dict[int, list[int]] = {0: [], 1: [], 2: [], 3: []}
    for m in rx.finditer(src):
        line_start = src.rfind("\n", 0, m.start()) + 1
        is_heading = (not src[line_start:m.start()].strip()
                      and src[m.end():m.end() + 2].lstrip()[:1] == ".")
        rest = src[m.start():]
        ends = [len(rest)]
        m_adopt = _SRC_ADOPTION_RE.search(rest, 30)
        if m_adopt:
            ends.append(m_adopt.end())
        for head_re in (_SRC_HEAD_OLD_RE, _SRC_HEAD_NEW_RE):
            m_head = head_re.search(rest, 60)
            if m_head:
                ends.append(m_head.start())
        span = rest[:min(ends)]
        n = _word_count(span)
        if not n:
            continue
        has_open = bool(_SOURCE_OPENING_RE.search(span))
        rank = 0 if (is_heading and has_open) else 1 if is_heading else 2 if has_open else 3
        ranked[rank].append(n)
    for rank in (0, 1, 2, 3):
        if ranked[rank]:
            return min(ranked[rank])
    return None


def source_conservation(symbol: str, fmt: str, archive_path: str | None,
                        converted_path: str | None, result: dict) -> dict:
    """Compare the parse against its archived source. Never consults extractor flags.

    status: 'ok' | 'truncated' | 'unchecked' | 'unreadable'
    """
    parse_words = parse_word_count(result)
    rep = {"status": "unchecked", "parse_words": parse_words,
           "source_words": None, "ratio": None, "reason": ""}
    if not (archive_path or converted_path):
        rep["reason"] = "no source file on the ledger row (volume-split child)"
        return rep
    try:
        src = read_source_text(fmt, archive_path, converted_path)
    except Exception as exc:
        rep["status"] = "unreadable"
        rep["reason"] = f"{type(exc).__name__}: {exc}"
        return rep
    whole_words = _word_count(src)
    rep["whole_source_words"] = whole_words
    rep["whole_ratio"] = round(parse_words / whole_words, 4) if whole_words else None
    src_words = source_span_words(symbol, src) if fmt == "pdf" else whole_words
    rep["source_words"] = src_words
    lost_opening = bool(_SOURCE_OPENING_RE.search(src)) and not _SOURCE_OPENING_RE.search(
        " ".join((el.get("text") or "") for el in result["elements"]))

    # Floor 2 (no anchoring needed): a document cannot be a twentieth of its own
    # source page. It is the FALLBACK, used only when the span denominator cannot be
    # trusted -- no anchor at all, or a span too short to be a document (the anchor
    # landed in the page's table of contents: `A/RES/32/1` anchors on a 14-word TOC
    # entry and its span check would pass at 0.71 while the parse holds 10 of the
    # page's 675 words). A trusted span is target-specific and always wins, so a
    # short resolution on a dense volume page is not failed for being 4% of it.
    span_trusted = src_words is not None and src_words >= MIN_TRUSTED_SPAN_WORDS
    if not span_trusted and whole_words:
        # How many documents does this source print? Its own headings say. The
        # expected share is the page average, and a document that holds less than
        # WHOLE_SOURCE_FLOOR of the average document on its own page is truncated.
        # (A fixed fraction of the page cannot work: a page may print 3 documents or
        # 40 -- `A/RES/1768(XVII)` is 951 words on a 37,700-word volume page and is
        # complete.)
        n_docs = max(1, len(_SRC_HEAD_OLD_RE.findall(src))
                     + len(_SRC_HEAD_NEW_RE.findall(src))
                     + len(_SRC_HEAD_RESOLUTION_RE.findall(src)))
        expected = whole_words / n_docs
        rep["source_words"] = round(expected)
        rep["ratio"] = round(parse_words / expected, 4) if expected else None
        if expected and parse_words < WHOLE_SOURCE_FLOOR * expected:
            rep["status"] = "truncated"
            rep["reason"] = (
                f"parse holds {parse_words} words; its source prints {n_docs} document(s) "
                f"in {whole_words} words, so this one should hold about {expected:.0f} "
                f"({100 * parse_words / expected:.1f}% of it, floor "
                f"{100 * WHOLE_SOURCE_FLOOR:.0f}%)"
                + ("; source has an opening formula, parse has none" if lost_opening else ""))
            return rep
    if src_words is None:
        rep["reason"] = ("no source anchor for this symbol shape; only the whole-source "
                         f"floor applied (ratio {rep['whole_ratio']})")
        return rep
    if not src_words:
        rep["reason"] = "source holds no words"
        return rep
    ratio = parse_words / src_words
    rep["ratio"] = round(ratio, 4)
    rep["status"] = "ok" if ratio >= TRUNCATION_FLOOR else "truncated"
    if rep["status"] == "truncated":
        rep["reason"] = (f"parse holds {parse_words} of {src_words} source words "
                         f"({100 * ratio:.1f}%), floor {100 * TRUNCATION_FLOOR:.0f}%"
                         + ("; source has an opening formula, parse has none"
                            if lost_opening else ""))
    return rep


def _check_accounting(result: dict, raw_rows: list[dict]) -> str | None:
    """Return an error string if the accounting invariant is violated, else None.

    Two halves, because the position half alone is satisfied by an element whose
    text has nothing to do with the rows it claims (the `A-FABRICATED-TEXT`
    control: replace every element's text with invented prose, leave positions
    untouched, and the old check stayed quiet):

      * POSITIONS -- every raw position is consumed exactly once, by an element or
        by a ledgered drop;
      * TEXT -- every element's word multiset is a SUBSET of the union of the raw
        rows it names in `positions[]`. The parser may drop, reorder or re-prefix
        words; it may never introduce one. The only synthetic strings it is allowed
        to emit are named here (the 'Vote record' placeholder and the vote dict's
        own keys), so an allowance cannot quietly widen.
    """
    all_positions = {r["position"] for r in raw_rows}
    consumed: list[int] = []
    for el in result["elements"]:
        # a split element (one raw row cut into title + body) re-uses its parent's
        # positions; they are counted once, on the element that owns them.
        if el.get("split_continuation"):
            continue
        consumed.extend(el["positions"])
    for d in result["dropped"]:
        consumed.append(d["position"])
    consumed_set = set(consumed)
    if len(consumed) != len(consumed_set):
        dupes = [p for p in consumed_set if consumed.count(p) > 1]
        return f"duplicate positions: {sorted(dupes)[:10]}"
    missing = all_positions - consumed_set
    extra = consumed_set - all_positions
    if missing:
        return f"unaccounted positions: {sorted(missing)[:10]}"
    if extra:
        return f"phantom positions: {sorted(extra)[:10]}"

    by_pos = {r["position"]: (r.get("text") or "") for r in raw_rows}
    for idx, el in enumerate(result["elements"]):
        source = " ".join(by_pos.get(p, "") for p in el.get("positions") or [])
        have = _word_bag(source)
        want = _word_bag((el.get("prefix") or "") + " " + (el.get("text") or ""))
        if el.get("type") == "vote_record":
            for key, lst in (el.get("vote") or {}).items():
                if not lst:
                    continue          # empty tally: its label never appeared in the raw
                for w, c in _word_bag(" ".join(lst) + " " + key.replace("_", " ")).items():
                    want[w] = want.get(w, 0) + c
            for w in ("vote", "record"):          # the 'Vote record' placeholder
                want.pop(w, None)
        invented = [w for w, c in want.items() if c > have.get(w, 0)]
        if invented:
            return (f"invented text in element {idx} (type={el.get('type')}, "
                    f"positions={el.get('positions')}): {sorted(invented)[:8]}")
    return None


# ---------------------------------------------------------------------------
# Self-test: NEGATIVE CONTROLS for the sem-v5 guards
#
# Doctrine: "a check that has never been shown to fail is absent, not passing."
# Every control below damages a real-shaped input and asserts the guard REJECTS it,
# and every damaging control is paired with an undamaged one that must stay quiet,
# so a guard that simply always fires cannot pass this suite.
# ---------------------------------------------------------------------------


def _row(pos: int, text: str, kind: str = "paragraph", **props) -> dict:
    return {"position": pos, "kind": kind, "text": text, "style_id": props.pop("style", None),
            "style_name": None, "numbering": None, "props": props or {},
            "table_cell": None, "hyperlinks": None, "footnote_ref": None}


def _self_test() -> int:
    failures: list[str] = []

    def check(cond: bool, msg: str) -> None:
        if not cond:
            failures.append(msg)

    def parse(texts: list[str], fmt: str = "pdf") -> dict:
        return parse_document("TEST", fmt, [_row(i, t) for i, t in enumerate(texts)])

    clean_body = [
        "Resolution 661 (1990)",
        "The Security Council,",
        "Recalling its resolution 660 (1990),",
        "1. Determines that Iraq has failed to comply;",
        "2. Decides to take the following measures;",
        "3. Decides that all States shall prevent:",
        "(a) The import into their territories of all commodities;",
        "(b) Any activities by their nationals which would promote;",
        "4. Decides that all States shall not make available funds;",
        "5. Calls upon all States to act strictly in accordance;",
    ]

    # --- POSITIVE control: an undamaged sequence keeps every marker -----------
    good = parse(clean_body)
    pfx = [e.get("prefix") for e in good["elements"] if e.get("prefix")]
    check(pfx == ["1.", "2.", "3.", "(a)", "(b)", "4.", "5."],
          f"positive control: clean 1..5 sequence lost markers -> {pfx}")
    check(not [i for i in good["issues"] if i["problem"] == "unconfirmed_marker"],
          "positive control: clean sequence produced marker refusals")

    # --- NEGATIVE control 1: the S/RES/661(1990) fabrication ------------------
    damaged = list(clean_body)
    damaged.insert(8, "19. nationals or in their territories which promote or are "
                      "calculated to promote such sale or supply;")
    bad = parse(damaged)
    prefixes = [e.get("prefix") for e in bad["elements"] if e.get("prefix")]
    check("19." not in prefixes,
          f"NEGATIVE CONTROL FAILED: invented marker '19.' was accepted -> {prefixes}")
    check([p for p in prefixes if p in ("3.", "4.", "5.")] == ["3.", "4.", "5."],
          f"invented-marker control damaged the genuine sequence -> {prefixes}")
    refusals = [i for i in bad["issues"] if i["problem"] == "unconfirmed_marker"]
    check(len(refusals) == 1 and refusals[0]["text_head"].startswith("19."),
          f"refused marker was not ledgered -> {refusals}")
    kept_text = " ".join(e.get("text") or "" for e in bad["elements"])
    check("19. nationals or in their territories" in kept_text,
          "refusing a marker DELETED text; it must stay inside the element")

    # --- NEGATIVE control 2: a page number ahead of the sequence --------------
    pagenum = list(clean_body)
    pagenum.insert(6, "26. which the Council will need to take further measures under;")
    pn = parse(pagenum)
    check("26." not in [e.get("prefix") for e in pn["elements"]],
          "NEGATIVE CONTROL FAILED: out-of-band number '26.' accepted as a marker")

    # --- NEGATIVE control 2b: page numbers that form their own tidy run -------
    # A table-of-contents page carries '59.', '60.', '61.' as PAGE numbers; they
    # chain perfectly, so only the band (a document cannot number more paragraphs
    # than it has rows) can reject them.
    toc = parse([
        "Resolution 1110 (XI)",
        "The General Assembly,",
        "59. Admission of Morocco to membership in the United Nations",
        "60. Admission of Tunisia to membership in the United Nations",
        "61. Report of the Security Council on its work",
    ])
    check(not [e.get("prefix") for e in toc["elements"] if e.get("prefix")],
          "NEGATIVE CONTROL FAILED: page numbers '59.'-'61.' accepted as markers "
          f"-> {[e.get('prefix') for e in toc['elements'] if e.get('prefix')]}")

    # --- NEGATIVE control 3: citation tail read as a Roman subparagraph -------
    cite = list(clean_body)
    cite.insert(7, "(IX) of 11 December 1954 and 910 (X) of 29 November 1955,")
    ct = parse(cite)
    check("(IX)" not in [e.get("prefix") for e in ct["elements"]],
          "NEGATIVE CONTROL FAILED: citation tail '(IX)' accepted as a subparagraph marker")
    roman = parse(clean_body[:6] + [
        "(I) To examine the reports on the progress of implementation;",
        "(II) To seek from all States further information;",
        "(III) To report on its work to the Council;",
    ])
    check([e.get("prefix") for e in roman["elements"] if e.get("prefix") and "I" in e["prefix"]]
          == ["(I)", "(II)", "(III)"],
          "positive control: a confirmed (I)(II)(III) run was refused")

    # --- NEGATIVE control 4: word accounting sees text vanish ----------------
    rows = [_row(i, t) for i, t in enumerate(clean_body)]
    res = parse_document("TEST", "pdf", rows)
    check(check_word_accounting(res, rows) == [],
          "positive control: word accounting reports loss on an undamaged parse")
    res["elements"][3]["text"] = res["elements"][3]["text"].replace("Iraq", "")
    lost = check_word_accounting(res, rows)
    check(lost == ["iraq"],
          f"NEGATIVE CONTROL FAILED: word accounting missed a deleted word -> {lost}")

    # --- NEGATIVE control 4b: A-FABRICATED-TEXT ------------------------------
    # Positions intact, text invented. The position-only invariant stayed quiet on
    # exactly this damage, which is why the parser could report "0 accounting
    # failures" while the corpus held invented markers.
    fab_rows = [_row(i, t) for i, t in enumerate(clean_body)]
    fab = parse_document("TEST", "pdf", fab_rows)
    check(_check_accounting(fab, fab_rows) is None,
          "positive control: accounting reports invented text on an undamaged parse")
    fab["elements"][4]["text"] = ("The Council authorizes the use of force against "
                                  "any State that fails to comply.")
    err_fab = _check_accounting(fab, fab_rows)
    check(err_fab is not None and "invented text" in err_fab,
          f"NEGATIVE CONTROL FAILED (A-FABRICATED-TEXT): invented prose accepted "
          f"-> {err_fab!r}")

    # --- NEGATIVE control 4c: decision bodies must not be frontmatter --------
    # A GA decision has no opening formula. Before sem-v5 the state machine never
    # left `front`, so the decision itself was typed 'frontmatter' and hidden.
    dec = parse([
        "77/408. Appointment of members of the Advisory Committee",
        "At its 87th plenary meeting, on 30 June 2023, the General Assembly, on the "
        "recommendation of the Fifth Committee, appointed Minhong Yi (Republic of "
        "Korea) as a member of the Advisory Committee.",
        "As a result, the Advisory Committee is composed as follows: Surendra Kumar "
        "Adhana (India), Yves Eric Ahoussougbemey (Benin).",
    ])
    types = [e["type"] for e in dec["elements"]]
    check("frontmatter" not in types,
          f"NEGATIVE CONTROL FAILED: a decision body was typed frontmatter -> {types}")
    check(types.count("paragraph") == 2,
          f"decision body did not become two paragraphs -> {types}")

    # --- NEGATIVE control 4d: title and decision on ONE line -----------------
    one_line = parse([
        "77/547. Problems arising from the accumulation of conventional ammunition "
        "stockpiles in surplus At its 56th (resumed) plenary meeting, on 30 December "
        "2022, the General Assembly decided to defer consideration of the item.",
    ])
    check([e["type"] for e in one_line["elements"]] == ["title", "paragraph"],
          "NEGATIVE CONTROL FAILED: title+decision stayed one element -> "
          f"{[(e['type'], (e['text'] or '')[:40]) for e in one_line['elements']]}")
    check(len(one_line["elements"]) == 2
          and one_line["elements"][0]["text"].endswith("surplus")
          and one_line["elements"][1]["text"].startswith("At its 56th"),
          "title/body split fell in the wrong place -> "
          f"{[(e['text'] or '')[:60] for e in one_line['elements']]}")
    split_rows = [_row(0, " ".join(e["text"] for e in one_line["elements"]))]
    check(_check_accounting(parse_document("TEST", "pdf", split_rows), split_rows) is None,
          "a split title/body row broke the accounting invariant")

    # --- NEGATIVE control 4e: the crop that runs PAST the document -----------
    over = {"elements": [
        {"type": "frontmatter", "text": "Resolution 1004 (ES-Il)", "prefix": None},
        {"type": "opening", "text": "The General Assembly,", "prefix": None},
        {"type": "paragraph", "text": "Considering that the United Nations is based "
                                      "on sovereign equality,", "prefix": None},
    ]}
    check(check_overreach("A/RES/1005(ES-II)", over) == ["1004"],
          "NEGATIVE CONTROL FAILED: a neighbouring resolution's heading inside the "
          f"parse was not detected -> {check_overreach('A/RES/1005(ES-II)', over)}")
    cite_only = {"elements": [
        {"type": "title", "text": "Right of peoples to self-determination", "prefix": "79/104."},
        {"type": "paragraph", "text": "Recalling its resolution 1514 (XV) of 14 "
                                      "December 1960,", "prefix": None},
        {"type": "footnote", "text": "Resolution 1514 (XV).", "prefix": None},
    ]}
    check(check_overreach("A/RES/79/104", cite_only) == [],
          "positive control: a citation of another resolution was mistaken for "
          f"over-reach -> {check_overreach('A/RES/79/104', cite_only)}")

    # --- NEGATIVE control 5: source conservation sees truncation -------------
    import tempfile
    src_body = ("701 (VII). Korea: reports of the United Nations Agent General\n"
                "The General Assembly,\n"
                + " ".join(f"word{i}" for i in range(300)) + "\n"
                "425th plenary meeting, 8 April 1953.\n"
                "702 (VII). Another decision entirely\n"
                + " ".join(f"other{i}" for i in range(300)) + "\n")
    with tempfile.NamedTemporaryFile("w", suffix=".txt", delete=False) as fh:
        fh.write(src_body)
        src_path = fh.name
    full = {"elements": [{"type": "opening", "text": "The General Assembly,", "prefix": None}]
            + [{"type": "paragraph", "text": " ".join(f"word{i}" for i in range(300)),
                "prefix": None}]}
    rep_ok = source_conservation("A/RES/701(VII)", "pdf", src_path, None, full)
    check(rep_ok["status"] == "ok",
          f"positive control: a complete parse was called {rep_ok['status']} ({rep_ok})")
    stub = {"elements": [{"type": "title", "text": "701 (VII). Korea: reports of the "
                                                   "United Nations", "prefix": None}]}
    rep_bad = source_conservation("A/RES/701(VII)", "pdf", src_path, None, stub)
    check(rep_bad["status"] == "truncated",
          f"NEGATIVE CONTROL FAILED: title-only parse was not called truncated ({rep_bad})")
    rep_missing = source_conservation("A/RES/701(VII)", "pdf",
                                      src_path + ".does-not-exist", None, full)
    check(rep_missing["status"] == "unreadable",
          f"NEGATIVE CONTROL FAILED: a missing source file did not fail ({rep_missing})")
    Path(src_path).unlink()

    # --- NEGATIVE control 5b: the anchor lands in a table of contents --------
    # `A/RES/32/1`'s only anchor is its TOC entry, so the span denominator is 14
    # words and the span check would pass at 0.71 on a 10-word parse. The span must
    # be distrusted below MIN_TRUSTED_SPAN_WORDS and the page-average floor decide.
    toc_src = ("701 (VII). Korea: reports of the Agent General ......... 13\n"
               "702 (VII). Financial reports and accounts ......... 14\n"
               "703 (VII). Question of the Trust Territory ......... 15\n"
               + " ".join(f"body{i}" for i in range(420)) + "\n")
    with tempfile.NamedTemporaryFile("w", suffix=".txt", delete=False) as fh:
        fh.write(toc_src)
        toc_path = fh.name
    toc_stub = {"elements": [{"type": "title", "text": "701 (VII). Korea: reports of "
                                                       "the Agent General", "prefix": None}]}
    rep_toc = source_conservation("A/RES/701(VII)", "pdf", toc_path, None, toc_stub)
    check(rep_toc["status"] == "truncated",
          f"NEGATIVE CONTROL FAILED: a parse of one TOC line passed because its span "
          f"anchor landed in the table of contents ({rep_toc})")
    toc_full = {"elements": [{"type": "paragraph",
                              "text": " ".join(f"body{i}" for i in range(140)),
                              "prefix": None}]}
    check(source_conservation("A/RES/701(VII)", "pdf", toc_path, None,
                              toc_full)["status"] == "ok",
          "positive control: a document holding its page share was called truncated")
    Path(toc_path).unlink()

    # --- NEGATIVE control 6: the span must not swallow the neighbour ---------
    span = source_span_words("A/RES/701(VII)", src_body)
    check(span is not None and 290 < span < 320,
          f"span anchoring wrong: expected ~305 words for the target only, got {span}")

    for f in failures:
        print(f"  FAIL {f}")
    print(f"\n{'FAIL' if failures else 'PASS'} — self-test: "
          f"{len(failures)} failing control(s)")
    return 1 if failures else 0


BATCH_DOCS = 20  # docs per short-lived DB connection (mirrors fulltext_extract_raw.py)


def main() -> int:
    ap = argparse.ArgumentParser(description="Semantic full-text parser (v1)")
    ap.add_argument("--limit", type=int)
    ap.add_argument("--offset", type=int, default=0)
    ap.add_argument("--symbol")
    ap.add_argument("--out", default=str(OUT_DIR))
    ap.add_argument("--to-db", action="store_true",
                    help="load the semantic DB tables (migration 003) alongside JSON")
    ap.add_argument("--db-only", action="store_true",
                    help="load the semantic DB only; skip writing JSON files")
    ap.add_argument("--self-test", action="store_true",
                    help="run the sem-v5 negative controls (no DB, no archive) and exit")
    args = ap.parse_args()

    if args.self_test:
        return _self_test()

    to_db = args.to_db or args.db_only
    write_json = not args.db_only

    out_dir = Path(args.out)
    if write_json:
        out_dir.mkdir(parents=True, exist_ok=True)

    targets = fetch_targets(args.limit, args.symbol, args.offset)
    dest = []
    if write_json:
        dest.append(str(out_dir))
    if to_db:
        dest.append("db(document_paragraphs, document_parses)")
    print(f"Parsing {len(targets)} documents -> {', '.join(dest) or '(nothing)'}")

    n_ok = 0
    n_acct_fail = 0
    n_word_fail = 0
    n_loaded = 0
    n_failed = 0
    n_truncated = 0
    n_unreadable = 0
    n_unchecked = 0
    n_overreach = 0
    n_markers_refused = 0
    total_elems = 0
    truncated_syms: list[str] = []
    for start in range(0, len(targets), BATCH_DOCS):
        chunk = targets[start:start + BATCH_DOCS]
        with get_conn() as conn:
            for symbol, lang, fmt, archive_path, converted_path in chunk:
                try:
                    raw_rows = fetch_rows(conn, symbol, lang)
                    result = parse_document(symbol, fmt, raw_rows)
                    n_markers_refused += sum(
                        1 for i in result["issues"] if i["problem"] == "unconfirmed_marker")
                    err = _check_accounting(result, raw_rows)
                    if err:
                        n_acct_fail += 1
                        result.setdefault("issues", []).append(
                            {"position": -1, "problem": "accounting", "text_head": err})
                        print(f"  ! {symbol}: ACCOUNTING {err}")
                    lost = check_word_accounting(result, raw_rows)
                    if lost:
                        n_word_fail += 1
                        result["issues"].append(
                            {"position": -1, "problem": "word_accounting",
                             "text_head": f"{len(lost)} unaccounted word type(s): "
                                          + ", ".join(lost[:12])})
                        print(f"  ! {symbol}: WORDS unaccounted {lost[:8]}")
                    # source conservation: the check the 145 silently truncated
                    # documents would have failed. Runs on every document; an
                    # unreadable or unanchorable source is reported, never a pass.
                    cons = source_conservation(symbol, fmt, archive_path,
                                               converted_path, result)
                    result["source_conservation"] = cons
                    # the other direction of the same crop defect: text belonging to
                    # the NEIGHBOURING document stored under this symbol.
                    foreign = check_overreach(symbol, result)
                    if foreign:
                        n_overreach += 1
                        result["issues"].append(
                            {"position": -1, "problem": "source_overreach",
                             "text_head": "headings of other documents inside this "
                                          f"parse: {', '.join(foreign)}"})
                        print(f"  ! {symbol}: OVERREACH carries headings of {foreign}")
                    if cons["status"] != "ok":
                        result["issues"].append(
                            {"position": -1, "problem": f"source_{cons['status']}",
                             "text_head": cons["reason"]
                                          or f"ratio={cons['ratio']}"})
                    if cons["status"] == "truncated":
                        n_truncated += 1
                        truncated_syms.append(symbol)
                        print(f"  ! {symbol}: TRUNCATED {cons['reason']}")
                    elif cons["status"] == "unreadable":
                        n_unreadable += 1
                        print(f"  ! {symbol}: SOURCE UNREADABLE {cons['reason']}")
                    elif cons["status"] == "unchecked":
                        n_unchecked += 1
                    if write_json:
                        out_path = out_dir / f"{sanitize_symbol(symbol)}.json"
                        out_path.write_text(
                            json.dumps(result, ensure_ascii=False, indent=1),
                            encoding="utf-8")
                    if to_db:
                        total_elems += load_document(conn, symbol, lang, fmt, result)
                        # A truncated / unverifiable-source parse is NEVER advanced to
                        # 'parsed': the rows are still written (nothing is deleted) but
                        # the ledger says the document is not fit, so the class cannot
                        # disappear again behind a green status.
                        if cons["status"] in ("truncated", "unreadable") or foreign:
                            reason = (f"source_{cons['status']}: {cons['reason']}"
                                      if cons["status"] in ("truncated", "unreadable")
                                      else f"source_overreach: carries headings of "
                                           f"{', '.join(foreign)}")
                            upsert_document_file(
                                conn, symbol, lang, status="parse_failed",
                                error=reason[:500])
                        else:
                            upsert_document_file(conn, symbol, lang, status="parsed", error=None)
                        conn.commit()
                        n_loaded += 1
                    n_ok += 1
                except Exception as exc:  # never crash the batch on one doc
                    if to_db:
                        conn.rollback()
                        try:
                            upsert_document_file(
                                conn, symbol, lang, status="parse_failed",
                                error=f"{type(exc).__name__}: {exc}"[:500])
                            conn.commit()
                        except Exception:
                            conn.rollback()
                    n_failed += 1
                    print(f"  ! {symbol}: {type(exc).__name__}: {exc}")
        done = start + len(chunk)
        if done % 100 == 0 or done == len(targets):
            print(f"  parsed {done}/{len(targets)} ok={n_ok} loaded={n_loaded} failed={n_failed}")

    print(f"\nDone: {n_ok} parsed, {n_acct_fail} accounting failures, "
          f"{n_word_fail} word-accounting failures, {n_failed} failed.")
    print(f"Markers refused as unconfirmed: {n_markers_refused}")
    print(f"Documents carrying another document's headings (crop over-reach): {n_overreach}")
    print(f"Source conservation: {n_truncated} TRUNCATED, {n_unreadable} unreadable, "
          f"{n_unchecked} unchecked, "
          f"{n_ok - n_truncated - n_unreadable - n_unchecked} ok")
    if truncated_syms:
        print("  truncated: " + ", ".join(truncated_syms[:20])
              + (" …" if len(truncated_syms) > 20 else ""))
    if to_db:
        print(f"Loaded {n_loaded} docs, {total_elems} element rows into the semantic DB.")
    verdict_fail = (n_failed or n_truncated or n_unreadable or n_acct_fail
                    or n_word_fail or n_overreach)
    print(("FAIL — " if verdict_fail else "PASS — ")
          + f"{n_ok} documents, {n_truncated} truncated, {n_overreach} carrying another "
            f"document's text, {n_unreadable} unreadable sources, "
            f"{n_acct_fail + n_word_fail} accounting failures, {n_failed} hard failures")
    # Exit code covers every failure class, accounting included (it used to reflect
    # hard crashes only, so a run with accounting failures still exited 0).
    return 0 if not (n_failed or n_truncated or n_unreadable or n_acct_fail or n_word_fail or n_overreach) else 1


if __name__ == "__main__":
    raise SystemExit(main())
